from __future__ import annotations

import copy
import hashlib
import json
import shutil
import sqlite3
import zipfile
from datetime import datetime, timedelta, timezone
from decimal import Decimal
from http import HTTPStatus
from pathlib import Path

from docx import Document
from openpyxl import Workbook, load_workbook
from PIL import Image, ImageDraw

from webapp.infonalia_webapp.justificaciones_baja.application.service import (
    JustificationApplicationService,
)
from webapp.infonalia_webapp.justificaciones_baja.documents import (
    build_document_payload,
    validate_excel,
    validate_word,
)
from webapp.infonalia_webapp.justificaciones_baja.documents.payload import (
    IdentificationInput,
    NarrativeInput,
    TransportDocumentInput,
)
from webapp.infonalia_webapp.justificaciones_baja.imports import preview_xlsx
from webapp.infonalia_webapp.justificaciones_baja.persistence import (
    JustificationRepository,
    ensure_justificaciones_baja_schema,
)
from webapp.infonalia_webapp.justificaciones_baja.snapshot import load_snapshot
from webapp.infonalia_webapp.tests.test_import_endpoints import load_app_module
from webapp.infonalia_webapp.tests.test_justificaciones_baja_api import (
    dispatch,
    isolated_api_environment,
    json_request,
    make_handler,
)


REPOSITORY_ROOT = Path(__file__).resolve().parents[3]
E2E_ROOT = REPOSITORY_ROOT / "tmp" / "justificaciones_baja_e2e"
USER_ID = "usuario-e2e"
PRODUCT_HEADERS = (
    "Producto",
    "Características",
    "Cantidad",
    "Precio oferta",
    "Importe oferta",
    "Precio coste",
    "Importe coste",
)


class ScriptedRandom:
    def __init__(self, values: list[int]) -> None:
        self._values = iter(values)

    def randint(self, minimum: int, maximum: int) -> int:
        value = next(self._values)
        if not minimum <= value <= maximum:
            raise AssertionError("La secuencia aleatoria E2E queda fuera de la horquilla.")
        return value


class AdvancingClock:
    def __init__(self) -> None:
        self._value = datetime(2026, 7, 14, 8, 0, tzinfo=timezone.utc)

    def __call__(self) -> datetime:
        value = self._value
        self._value += timedelta(seconds=1)
        return value


def _prepare_workspace() -> dict[str, Path]:
    allowed_parent = (REPOSITORY_ROOT / "tmp").resolve(strict=False)
    root = E2E_ROOT.resolve(strict=False)
    root.relative_to(allowed_parent)
    if root.exists():
        shutil.rmtree(root)
    paths = {
        "root": root,
        "inputs": root / "inputs",
        "dropbox": root / "dropbox_ficticio",
        "downloads": root / "descargas_verificadas",
        "runtime": root / "runtime",
        "database": root / "justificaciones_e2e.sqlite3",
        "manifest": root / "manifest.json",
    }
    for key in ("inputs", "dropbox", "downloads", "runtime"):
        paths[key].mkdir(parents=True, exist_ok=True)
    return paths


def _database(path: Path) -> sqlite3.Connection:
    connection = sqlite3.connect(path)
    connection.row_factory = sqlite3.Row
    connection.execute("PRAGMA foreign_keys = ON")
    connection.execute(
        """
        CREATE TABLE licitaciones (
            id INTEGER PRIMARY KEY,
            expediente TEXT NOT NULL,
            objeto TEXT NOT NULL,
            organismo TEXT NOT NULL,
            ruta_carpeta TEXT NOT NULL DEFAULT ''
        )
        """
    )
    connection.execute(
        """
        CREATE TABLE clientes (
            id INTEGER PRIMARY KEY,
            razon_social TEXT NOT NULL
        )
        """
    )
    connection.execute(
        """
        INSERT INTO licitaciones (id, expediente, objeto, organismo, ruta_carpeta)
        VALUES (1, 'EXP-E2E-0001', 'Suministro ficticio para prueba E2E',
                'Órgano de contratación de ensayo', '2026/EXP-E2E-0001')
        """
    )
    connection.execute(
        "INSERT INTO clientes (id, razon_social) VALUES (1, 'Empresa Anónima de Ensayo, S.L.')"
    )
    ensure_justificaciones_baja_schema(connection)
    connection.commit()
    return connection


def _source_workbook(path: Path) -> None:
    workbook = Workbook()
    workbook.properties.creator = "Prueba automatizada"
    workbook.properties.lastModifiedBy = "Prueba automatizada"
    sheet = workbook.active
    sheet.title = "Lote 1"
    sheet.append(["Producto", "Características", "Cantidad", "Precio", "Importe"])
    sheet.append(["Producto estándar", "Formato A", 10, 20, 200])
    sheet.append(["Producto estándar", "Formato B", 20, 15, 300])
    sheet.append(["TOTAL", "", "", "", 500])
    workbook.save(path)
    workbook.close()


def _route_image(path: Path) -> bytes:
    image = Image.new("RGB", (640, 360), color=(245, 248, 252))
    drawing = ImageDraw.Draw(image)
    drawing.line(
        [(70, 280), (220, 90), (410, 135), (570, 280), (70, 280)],
        fill=(45, 92, 145),
        width=8,
    )
    for x, y in ((70, 280), (220, 90), (410, 135), (570, 280)):
        drawing.ellipse((x - 11, y - 11, x + 11, y + 11), fill=(181, 54, 19))
    drawing.text((70, 310), "RUTA FICTICIA - PRUEBA E2E", fill=(32, 45, 62))
    image.save(path, format="PNG")
    image.close()
    return path.read_bytes()


def _draft_with_import(created: dict, imported: dict) -> dict:
    draft = copy.deepcopy(created["draft"])
    draft["identification"].update(
        {
            "duration_description": "Diez semanas de ejecución",
            "place": "Municipio de ensayo",
            "date_text": "14 de julio de 2026",
        }
    )
    draft["products"] = imported["products"]
    draft["source"]["product_import"] = {
        "format": imported["format"],
        "source_sha256": imported["source_sha256"],
        "sheet": imported["sheet"],
        "mapping": imported["mapping"],
        "ignored_rows": imported["ignored_rows"],
    }
    draft["transport"].update(
        {
            "operational_weeks": 10,
            "weekly_deliveries": 1,
            "circular_kilometres": "50",
            "effective_decimal_hours": "2.5",
            "kilometre_rate": "0.6776",
            "hourly_rate": "39.75",
            "contract_stops": 1,
            "shared_orders": 15,
            "route_duration_text": "2 h 30 min",
        }
    )
    draft["financial"].update(
        {
            "declared_lot_offer": "500",
            "general_expense_base": "500",
            "general_expense_percentage": "0.05",
            "indirect_costs": "3",
        }
    )
    draft["transport_document"].update(
        {
            "observatory_date": "Abril de 2026",
            "observatory_url": "https://example.invalid/observatorio-ficticio.pdf",
        }
    )
    draft["narrative"].update(
        {
            "exposition": "Estimación económica ficticia sometida a validación del cliente.",
            "arguments": ["La distribución se integra en rutas logísticas ya consolidadas."],
            "acquisition_text": "Los costes unitarios son estimaciones revisables por el cliente.",
            "transport_text": "La ruta se prorratea entre pedidos compartidos.",
            "structure_text": "La estructura ordinaria absorbe la gestión del contrato.",
            "conclusion": "La oferta ficticia presenta un resultado económico positivo.",
        }
    )
    return draft


def _sha256(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest().upper()


def _canonical_hash(value: object) -> str:
    encoded = json.dumps(
        value,
        ensure_ascii=False,
        sort_keys=True,
        separators=(",", ":"),
    ).encode("utf-8")
    return hashlib.sha256(encoded).hexdigest().upper()


def _assert_anonymous_archive(path: Path) -> None:
    forbidden = (
        b"llangon",
        b"salvador",
        b"c:\\users",
        b"c:/users",
    )
    with zipfile.ZipFile(path) as archive:
        package_content = b"\n".join(
            archive.read(name).lower()
            for name in archive.namelist()
            if name.endswith((".xml", ".rels"))
        )
    assert not any(marker in package_content for marker in forbidden)


def _payload_from_frozen_version(version: dict, *, route_image_path: Path):
    loaded = load_snapshot(version["snapshot_json"])
    assert loaded.is_valid and loaded.snapshot is not None
    context = version["document_context"]
    narrative = context["narrative"]
    payload = build_document_payload(
        loaded.snapshot,
        IdentificationInput(**context["identification"]),
        NarrativeInput(
            **{
                **narrative,
                "arguments": tuple(narrative.get("arguments", ())),
                "pending_validation_fields": tuple(
                    narrative.get("pending_validation_fields", ())
                ),
            }
        ),
        TransportDocumentInput(**context["transport_document"]),
        route_image_path=route_image_path,
        route_image_logical_name=context["route_image"]["logical_name"],
        generated_at=context["generated_at"],
        generated_by=context["generated_by"],
        draft=True,
    )
    assert payload.control.snapshot_sha256 == version["snapshot_sha256"]
    return payload


def _download_documents(
    documents: list[dict],
    *,
    dropbox_base: Path,
    download_directory: Path,
    label: str,
) -> dict[str, Path]:
    result: dict[str, Path] = {}
    base = dropbox_base.resolve(strict=True)
    for document in documents:
        source = (base / Path(document["relative_path"])).resolve(strict=True)
        source.relative_to(base)
        assert source.name == document["file_name"]
        assert source.stat().st_size == document["size_bytes"]
        assert _sha256(source) == document["sha256"]
        destination = download_directory / f"{label}_{document['file_name']}"
        shutil.copyfile(source, destination)
        assert _sha256(destination) == document["sha256"]
        assert destination.read_bytes() == source.read_bytes()
        result[document["document_type"]] = destination
    return result


def _verify_excel(
    path: Path,
    *,
    snapshot_sha256: str,
    payload_sha256: str,
    line_ids: list[str],
) -> None:
    workbook = load_workbook(path, data_only=False, keep_links=False)
    try:
        assert workbook.sheetnames == [
            "Identificación",
            "Productos",
            "Transporte",
            "Resumen",
            "Auditoría",
        ]
        formulas = [
            cell.coordinate
            for sheet in workbook.worksheets
            for row in sheet.iter_rows()
            for cell in row
            if cell.data_type == "f"
            or (isinstance(cell.value, str) and cell.value.lstrip().startswith("="))
        ]
        assert formulas == []
        audit = {
            str(row[0].value): row[1].value
            for row in workbook["Auditoría"].iter_rows(min_row=3, max_col=2)
            if row[0].value is not None
        }
        assert audit["snapshot_sha256"] == snapshot_sha256
        assert audit["payload_sha256"] == payload_sha256
        assert audit["product_count"] == len(line_ids)

        products = workbook["Productos"]
        assert [products.cell(row, 1).value for row in range(3, 5)] == line_ids
        assert products.cell(3, 8).value == 13.5
        assert products.cell(3, 12).value == "manual"
        assert products.cell(3, 13).value is False
        assert products.cell(4, 12).value == "generado"
        assert products.cell(4, 13).value is True
        assert not getattr(workbook, "_external_links", ())
    finally:
        workbook.close()
    _assert_anonymous_archive(path)


def _verify_word(path: Path, *, line_ids: list[str]) -> None:
    document = Document(path)
    product_table = next(
        table
        for table in document.tables
        if tuple(cell.text.strip() for cell in table.rows[0].cells) == PRODUCT_HEADERS
    )
    assert len(product_table.rows) == len(line_ids) + 1
    all_text = "\n".join(
        [paragraph.text for paragraph in document.paragraphs]
        + [cell.text for table in document.tables for row in table.rows for cell in row.cells]
    )
    assert "BORRADOR ESTIMATIVO PENDIENTE DE VALIDACIÓN DEL CLIENTE" in all_text
    assert "Producto estándar" in all_text

    with zipfile.ZipFile(path) as archive:
        story_names = [
            name
            for name in archive.namelist()
            if name == "word/document.xml"
            or name.startswith("word/header")
            or name.startswith("word/footer")
        ]
        story_xml = b"".join(archive.read(name) for name in story_names)
        all_xml = b"".join(
            archive.read(name) for name in archive.namelist() if name.endswith(".xml")
        )
    assert b"{{" not in all_xml and b"{%" not in all_xml
    assert b"instrText" not in story_xml and b"fldSimple" not in story_xml
    assert b"LLangon" not in all_xml and b"Salvador" not in all_xml
    _assert_anonymous_archive(path)


def test_e2e_http_permissions_and_csrf() -> None:
    app = load_app_module()
    with isolated_api_environment(app) as environment:
        missing_csrf = make_handler(
            app,
            "POST",
            "/api/justificaciones-baja",
            {
                "licitacion_id": environment.licitacion_id,
                "cliente_id": environment.cliente_id,
            },
            csrf_token=None,
        )
        dispatch(missing_csrf, "POST")
        assert missing_csrf.responses[-1][0] == HTTPStatus.FORBIDDEN

        status, created, _ = json_request(
            environment,
            "POST",
            "/api/justificaciones-baja",
            {
                "licitacion_id": environment.licitacion_id,
                "cliente_id": environment.cliente_id,
                "lote_numero": "1",
                "lote_nombre": "Lote ficticio E2E",
                "importe_ofertado": "500",
            },
        )
        assert status == HTTPStatus.CREATED
        justification_id = int(created["item"]["id"])

        status, reviewer_listing, _ = json_request(
            environment,
            "GET",
            "/api/justificaciones-baja",
            role="nuria",
            username="nuria_e2e",
        )
        assert status == HTTPStatus.OK
        assert reviewer_listing["permissions"]["view"] is True
        assert reviewer_listing["permissions"]["download"] is True
        assert reviewer_listing["permissions"]["edit"] is False

        status, forbidden, _ = json_request(
            environment,
            "POST",
            f"/api/justificaciones-baja/{justification_id}/costes/generar",
            {"revision": created["item"]["revision"]},
            role="nuria",
            username="nuria_e2e",
        )
        assert status == HTTPStatus.FORBIDDEN
        assert "permiso" in str(forbidden["error"]).lower()

        status, admin_listing, _ = json_request(
            environment,
            "GET",
            "/api/justificaciones-baja",
        )
        assert status == HTTPStatus.OK
        assert admin_listing["permissions"]["create"] is True
        assert admin_listing["permissions"]["generate_documents"] is True


def test_full_anonymized_justification_flow_generates_versioned_documents() -> None:
    paths = _prepare_workspace()
    source_workbook = paths["inputs"] / "oferta_anonima.xlsx"
    route_image = paths["inputs"] / "ruta_ficticia.png"
    _source_workbook(source_workbook)
    route_bytes = _route_image(route_image)

    connection = _database(paths["database"])
    repository = JustificationRepository(connection)
    service = JustificationApplicationService(
        repository,
        clock=AdvancingClock(),
        random_source=ScriptedRandom([40, 41, 45]),
        temporary_root=paths["runtime"],
    )
    try:
        created = service.create(
            licitacion={
                "id": 1,
                "expediente": "EXP-E2E-0001",
                "objeto": "Suministro ficticio para prueba E2E",
                "organismo": "Órgano de contratación de ensayo",
            },
            cliente={
                "id": 1,
                "razon_social": "Empresa Anónima de Ensayo, S.L.",
                "nif_cif": "B00000000",
                "domicilio_fiscal": "Calle Ficticia 1",
                "municipio": "Municipio de ensayo",
                "telefono_principal": "+34 000 000 000",
                "email_principal": "pruebas@example.invalid",
                "representante_nombre": "Persona de Ensayo",
                "representante_nif": "00000000T",
                "representante_cargo": "Representante",
            },
            lote_numero="1",
            lote_nombre="Productos de ensayo",
            declared_offer="500",
            user_id=USER_ID,
        )

        mapping = {
            "name": "A",
            "characteristics": "B",
            "quantity": "C",
            "offered_unit_price": "D",
            "offered_amount": "E",
        }
        imported = preview_xlsx(
            source_workbook.read_bytes(),
            filename=source_workbook.name,
            sheet_name="Lote 1",
            start_row=2,
            mapping=mapping,
        )
        repeated_preview = preview_xlsx(
            source_workbook.read_bytes(),
            filename=source_workbook.name,
            sheet_name="Lote 1",
            start_row=2,
            mapping=mapping,
        )
        assert imported["can_confirm"] is True
        assert imported["issues"] == []
        assert imported["ignored_rows"] == [4]
        assert len(imported["products"]) == 2
        assert [item["line_id"] for item in imported["products"]] == [
            item["line_id"] for item in repeated_preview["products"]
        ]

        saved = service.save(
            created["id"],
            draft=_draft_with_import(created, imported),
            expected_revision=created["revision"],
            user_id=USER_ID,
            event_type="products_imported",
            event_message="Oferta XLSX ficticia importada.",
            event_metadata={
                "source_sha256": imported["source_sha256"],
                "product_count": len(imported["products"]),
            },
        )
        assert all(product["generated_unit_cost"] is None for product in saved["draft"]["products"])

        generated = service.generate_costs(
            created["id"], expected_revision=saved["revision"], user_id=USER_ID
        )
        line_ids = [product["line_id"] for product in generated["draft"]["products"]]
        initial = {
            product["line_id"]: copy.deepcopy(product)
            for product in generated["draft"]["products"]
        }
        assert [
            product["applied_percentage"] for product in generated["draft"]["products"]
        ] == [40, 41]

        manual = service.set_manual_cost(
            created["id"],
            expected_revision=generated["revision"],
            line_id=line_ids[0],
            manual_unit_cost="13,500",
            user_id=USER_ID,
        )
        locked = service.set_product_lock(
            created["id"],
            expected_revision=manual["revision"],
            line_id=line_ids[1],
            locked=True,
            user_id=USER_ID,
        )
        recalculated = service.recalculate_costs(
            created["id"],
            expected_revision=locked["revision"],
            line_ids=None,
            user_id=USER_ID,
        )
        after = {product["line_id"]: product for product in recalculated["draft"]["products"]}
        assert after[line_ids[0]]["manual_unit_cost"] == "13.500"
        assert after[line_ids[0]]["cost_origin"] == "manual"
        assert after[line_ids[0]]["applied_percentage"] == 45
        assert (
            after[line_ids[0]]["generated_unit_cost"]
            != initial[line_ids[0]]["generated_unit_cost"]
        )
        assert after[line_ids[1]] == {
            **initial[line_ids[1]],
            "locked": True,
        }

        attached = service.attach_route_image(
            created["id"],
            expected_revision=recalculated["revision"],
            filename=route_image.name,
            content=route_bytes,
            user_id=USER_ID,
        )
        assert attached["draft"]["route_image"]["sha256"] == _sha256(route_image)
        assert attached["calculation"]["errors"] == []

        first_frozen = service.freeze(
            created["id"], expected_revision=attached["revision"], user_id=USER_ID
        )
        first_version = first_frozen["version"]
        assert first_version["version_number"] == 1
        assert first_version["snapshot_sha256"] == hashlib.sha256(
            first_version["snapshot_json"].encode("utf-8")
        ).hexdigest().upper()

        output = (
            paths["dropbox"]
            / "2026"
            / "EXP-E2E-0001"
            / "Justificaciones de baja"
            / "Lote_1"
        )
        first_generation = service.generate_documents(
            created["id"],
            version_number=1,
            output_directory=output,
            dropbox_base=paths["dropbox"],
            user_id=USER_ID,
        )
        assert first_generation["generation_number"] == 1
        assert len(first_generation["documents"]) == 2
        assert {document["document_type"] for document in first_generation["documents"]} == {
            "word",
            "excel",
        }
        assert {document["payload_sha256"] for document in first_generation["documents"]} == {
            first_generation["payload_sha256"]
        }
        first_downloads = _download_documents(
            first_generation["documents"],
            dropbox_base=paths["dropbox"],
            download_directory=paths["downloads"],
            label="version_001_generacion_001",
        )
        _verify_excel(
            first_downloads["excel"],
            snapshot_sha256=first_version["snapshot_sha256"],
            payload_sha256=first_generation["payload_sha256"],
            line_ids=line_ids,
        )
        _verify_word(first_downloads["word"], line_ids=line_ids)
        first_payload = _payload_from_frozen_version(
            first_version,
            route_image_path=route_image,
        )
        assert first_payload.sha256 == first_generation["payload_sha256"]
        assert Decimal(first_payload.summary.profit.raw) > Decimal("0")
        assert "positivo" in first_payload.narrative.conclusion.lower()
        assert validate_excel(first_downloads["excel"], first_payload).is_valid
        assert validate_word(first_downloads["word"], first_payload).is_valid

        first_file_hashes = {
            document["relative_path"]: document["sha256"]
            for document in first_generation["documents"]
        }
        before_reopen = service.get(created["id"])
        draft_hash = _canonical_hash(before_reopen["draft"])
        revision = before_reopen["revision"]
        reopened = service.get(created["id"])
        assert reopened["revision"] == revision
        assert _canonical_hash(reopened["draft"]) == draft_hash
        assert reopened["draft_frozen"] is True
        assert len(reopened["documents"]) == 2

        second_draft = copy.deepcopy(reopened["draft"])
        second_draft["narrative"]["conclusion"] = (
            "La segunda versión ficticia mantiene la viabilidad y amplía la conclusión."
        )
        second_saved = service.save(
            created["id"],
            draft=second_draft,
            expected_revision=reopened["revision"],
            user_id=USER_ID,
            event_type="second_version_draft",
            event_message="Borrador ficticio preparado para la segunda versión.",
        )
        assert second_saved["draft_based_on_version"] == 1
        assert second_saved["draft_frozen"] is False
        second_frozen = service.freeze(
            created["id"], expected_revision=second_saved["revision"], user_id=USER_ID
        )
        assert second_frozen["version"]["version_number"] == 2

        second_generation = service.generate_documents(
            created["id"],
            version_number=2,
            output_directory=output,
            dropbox_base=paths["dropbox"],
            user_id=USER_ID,
        )
        repeated_second_generation = service.generate_documents(
            created["id"],
            version_number=2,
            output_directory=output,
            dropbox_base=paths["dropbox"],
            user_id=USER_ID,
        )
        assert second_generation["generation_number"] == 1
        assert repeated_second_generation["generation_number"] == 2
        assert second_generation["payload_sha256"] == repeated_second_generation["payload_sha256"]
        all_second_paths = {
            document["relative_path"]
            for result in (second_generation, repeated_second_generation)
            for document in result["documents"]
        }
        assert len(all_second_paths) == 4
        assert not (set(first_file_hashes) & all_second_paths)

        for relative_path, expected_hash in first_file_hashes.items():
            assert _sha256(paths["dropbox"] / Path(relative_path)) == expected_hash

        final_item = service.get(created["id"])
        assert [version["version_number"] for version in final_item["versions"]] == [2, 1]
        assert len(final_item["documents"]) == 6
        assert len({document["relative_path"] for document in final_item["documents"]}) == 6
        history_types = [entry["event_type"] for entry in final_item["history"]]
        assert {
            "created",
            "products_imported",
            "costs_generated",
            "manual_cost_set",
            "product_locked",
            "costs_recalculated",
            "route_image_attached",
            "new_draft_from_version",
            "second_version_draft",
            "frozen",
            "document_word",
            "document_excel",
        } <= set(history_types)
        assert {entry["created_by"] for entry in final_item["history"]} == {USER_ID}

        resolved_document = repository.get_document_by_id(final_item["documents"][0]["id"])
        assert resolved_document["justificacion_id"] == created["id"]
        assert resolved_document["cliente_id"] == 1
        assert resolved_document["licitacion_id"] == 1

        database_dump = "\n".join(connection.iterdump())
        assert "LLangon" not in database_dump and "Salvador" not in database_dump

        manifest = {
            "kind": "JUSTIFICACIONES_BAJA_E2E_ANONIMIZADO",
            "source_xlsx_sha256": _sha256(source_workbook),
            "route_image_sha256": _sha256(route_image),
            "versions": [
                {
                    "version_number": 1,
                    "snapshot_sha256": first_version["snapshot_sha256"],
                    "payload_sha256": first_generation["payload_sha256"],
                },
                {
                    "version_number": 2,
                    "snapshot_sha256": second_frozen["version"]["snapshot_sha256"],
                    "payload_sha256": second_generation["payload_sha256"],
                },
            ],
            "documents": [
                {
                    "version_number": document["version_number"],
                    "generation_number": document["generation_number"],
                    "document_type": document["document_type"],
                    "relative_path": document["relative_path"],
                    "sha256": document["sha256"],
                    "payload_sha256": document["payload_sha256"],
                }
                for document in final_item["documents"]
            ],
            "history_event_count": len(final_item["history"]),
            "authorization": {
                "http_e2e_test": "test_e2e_http_permissions_and_csrf",
                "csrf_required_for_mutations": True,
                "admin_can_mutate": True,
                "reviewer_can_view_and_download": True,
                "reviewer_cannot_mutate": True,
            },
        }
        paths["manifest"].write_text(
            json.dumps(manifest, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )
        assert paths["manifest"].is_file()
        manifest_text = paths["manifest"].read_text(encoding="utf-8").lower()
        assert not any(
            marker in manifest_text
            for marker in ("llangon", "salvador", "c:\\users", "c:/users")
        )
    finally:
        connection.close()
