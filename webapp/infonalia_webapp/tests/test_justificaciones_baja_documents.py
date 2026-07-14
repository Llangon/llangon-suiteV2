from __future__ import annotations

import ast
import hashlib
import re
import zipfile
from dataclasses import replace
from decimal import Decimal
from pathlib import Path
from xml.etree import ElementTree as ET

import pytest
from docx import Document
from docx.oxml.ns import qn
from openpyxl import load_workbook
from PIL import Image

from webapp.infonalia_webapp.justificaciones_baja import (
    CostOrigin,
    SnapshotMetadata,
    calculate_justification,
    create_snapshot,
)
from webapp.infonalia_webapp.justificaciones_baja.documents import (
    build_document_payload,
    generate_excel,
    generate_word,
    validate_excel,
    validate_word,
)
from webapp.infonalia_webapp.justificaciones_baja.documents.filenames import (
    UnsafeDocumentPathError,
    next_versioned_path,
    publish_atomic_no_overwrite,
    safe_component,
    temporary_output_path,
)
from webapp.infonalia_webapp.justificaciones_baja.documents.payload import (
    IdentificationInput,
    NarrativeInput,
    RawDisplayValue,
    TransportDocumentInput,
)
from webapp.infonalia_webapp.justificaciones_baja.documents.template_manifest import (
    DEFAULT_TEMPLATE_PATH,
    WORD_TEMPLATE_SENTINEL,
)
from webapp.infonalia_webapp.justificaciones_baja.documents.validators import (
    DocumentValidationError,
    InvalidRouteImageError,
    inspect_route_image,
    validate_payload,
    validate_template,
)
from webapp.infonalia_webapp.tests.fixtures.justificaciones_baja_salvador import (
    salvador_justification,
)


REPO_ROOT = Path(__file__).resolve().parents[3]
REFERENCE_DIR = REPO_ROOT / "docs" / "justificaciones_baja" / "referencia"
REFERENCE_HASHES = {
    "Plantilla para bajas.docx": "B48A4EE3369371BC36A88E271EDE3B9E0EDD220A36CD9BA99C8AE09B55C9B155",
    "Plantilla Excel para bajas - LOTE 1.xlsx": "F9B756CFA693ACE97C72A268DFE8CA57692443D89F0D8DE093E65A24085ACC19",
    "Baja Lote 1 Salvador.pdf": "A7805BA80A6C081D5AF7AB8F758AF88977C0D8781A06C781FB35EADFFEBB60BD",
}


def _sha256(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest().upper()


def _route_image(path: Path, image_format: str = "PNG", size: tuple[int, int] = (640, 320)) -> Path:
    image = Image.new("RGB", size, "#FCE7DF")
    image.save(path, format=image_format)
    return path


def _snapshot(product_count: int = 26, *, mode: str = "mixed", locked: bool = False, long: bool = False):
    draft = salvador_justification()
    source = draft.products
    products = []
    for index in range(product_count):
        original = source[index % len(source)]
        line_id = original.line_id if product_count <= len(source) else f"PX{index + 1:03d}"
        name = original.name
        characteristics = original.characteristics
        if long:
            characteristics = (characteristics + " · descripción extensa para comprobar ajuste de línea") * 8
        if mode == "generated":
            manual = None
            origin = CostOrigin.GENERATED
        elif mode == "manual":
            manual = original.manual_unit_cost or original.generated_unit_cost
            origin = CostOrigin.MANUAL
        else:
            manual = original.manual_unit_cost
            origin = original.cost_origin
        products.append(
            replace(
                original,
                line_id=line_id,
                name=name,
                characteristics=characteristics,
                manual_unit_cost=manual,
                cost_origin=origin,
                locked=locked,
            )
        )
    draft = replace(draft, products=tuple(products))
    calculation = calculate_justification(draft)
    assert calculation.is_valid, calculation.errors
    built = create_snapshot(
        draft,
        calculation,
        SnapshotMetadata(created_at="2026-07-14T12:00:00+00:00", created_by="test-suite"),
    )
    assert built.snapshot is not None, built.errors
    return built.snapshot


def _payload(
    tmp_path: Path,
    *,
    product_count: int = 26,
    mode: str = "mixed",
    locked: bool = False,
    long: bool = False,
    image_format: str | None = None,
    identification: IdentificationInput | None = None,
):
    image_path = None
    if image_format:
        suffix = ".jpg" if image_format == "JPEG" else ".png"
        image_path = _route_image(tmp_path / f"route{suffix}", image_format)
    identity = identification or IdentificationInput(
        expediente="EXP-ANON-2026-001",
        organismo="Organismo público de prueba",
        objeto="Suministro alimentario anonimizado",
        lot_number="1",
        lot_name="Panadería y bollería",
        duration_description="24 meses",
        place="Granada",
        date_text="14 de julio de 2026",
        client="EMPRESA ALIMENTARIA DE PRUEBA, S.L.",
        nif="B00000000",
        address="Domicilio de prueba",
        phone="000 000 000",
        email="pruebas@example.invalid",
        representative="Representante de prueba",
        representative_dni="00000000T",
        role="Representante",
        signatory="Firmante de prueba",
    )
    payload = build_document_payload(
        _snapshot(product_count, mode=mode, locked=locked, long=long),
        identity,
        NarrativeInput(
            subject="Justificación de la viabilidad económica de la oferta.",
            exposition="Se presenta un borrador estimativo sujeto a validación del cliente.",
            arguments=("Operativa habitual disponible.", "Agrupación logística de pedidos."),
            acquisition_text="Los costes de adquisición proceden del snapshot congelado.",
            transport_text="El transporte utiliza los parámetros confirmados.",
            structure_text="Se imputa la estructura prevista en el cálculo.",
            conclusion="El borrador arroja un resultado económico positivo.",
            pending_validation_fields=("costes unitarios",),
        ),
        TransportDocumentInput(
            observatory="Observatorio de Costes del Transporte",
            observatory_date="abril de 2026",
            observatory_url="https://www.transportes.gob.es/",
            vehicle="Vehículo rígido de dos ejes de distribución",
        ),
        route_image_path=image_path,
        route_image_logical_name="ruta_circular" + (image_path.suffix if image_path else ""),
        generated_at="2026-07-14T12:00:00+00:00",
        generated_by="test-suite",
    )
    return payload, image_path


def _rewrite_zip(source: Path, destination: Path, replacements: dict[str, bytes], additions: dict[str, bytes] | None = None) -> None:
    with zipfile.ZipFile(source) as original, zipfile.ZipFile(destination, "w", zipfile.ZIP_DEFLATED) as output:
        for item in original.infolist():
            output.writestr(item, replacements.get(item.filename, original.read(item.filename)))
        for name, data in (additions or {}).items():
            output.writestr(name, data)


def _word_semantics(path: Path) -> tuple[tuple[str, ...], tuple[tuple[tuple[str, ...], ...], ...]]:
    document = Document(path)
    paragraphs = tuple(p.text for p in document.paragraphs)
    tables = tuple(
        tuple(tuple(cell.text for cell in row.cells) for row in table.rows)
        for table in document.tables
    )
    return paragraphs, tables


def _excel_semantics(path: Path) -> dict[str, tuple[tuple[object, ...], ...]]:
    workbook = load_workbook(path, data_only=False)
    result = {
        sheet.title: tuple(tuple(cell.value for cell in row) for row in sheet.iter_rows())
        for sheet in workbook.worksheets
    }
    workbook.close()
    return result


def test_t01_caso_salvador_completo(tmp_path: Path) -> None:
    payload, image = _payload(tmp_path, image_format="PNG")
    word = generate_word(payload, tmp_path.resolve(), route_image_path=image)
    excel = generate_excel(payload, tmp_path.resolve(), route_image_path=image)
    assert len(payload.products) == 26
    assert word.snapshot_sha256 == excel.snapshot_sha256 == payload.control.snapshot_sha256
    assert word.payload_sha256 == excel.payload_sha256 == payload.sha256
    assert validate_word(word.path, payload).is_valid
    assert validate_excel(excel.path, payload).is_valid


@pytest.mark.parametrize(
    ("case", "count"),
    (("T03", 1), ("T04", 5), ("T05", 26), ("T06", 60)),
)
def test_t02_to_t06_product_counts(case: str, count: int, tmp_path: Path) -> None:
    payload, _ = _payload(tmp_path, product_count=count)
    word = generate_word(payload, tmp_path.resolve())
    excel = generate_excel(payload, tmp_path.resolve())
    assert len(payload.products) == count, case
    assert validate_word(word.path, payload).is_valid
    assert validate_excel(excel.path, payload).is_valid


def test_t02_lote_sin_productos_es_rechazado_por_fase_1() -> None:
    draft = replace(salvador_justification(), products=())
    calculation = calculate_justification(draft)
    assert not calculation.is_valid
    assert "sin_productos" in {issue.code for issue in calculation.errors}


def test_t07_descripciones_largas(tmp_path: Path) -> None:
    payload, _ = _payload(tmp_path, product_count=5, long=True)
    result = generate_word(payload, tmp_path.resolve())
    assert validate_word(result.path, payload).is_valid


def test_t08_nombres_duplicados_con_line_id_distinto(tmp_path: Path) -> None:
    payload, _ = _payload(tmp_path, product_count=5)
    assert len({item.line_id for item in payload.products}) == 5
    assert len({item.name for item in payload.products}) < 5
    assert validate_payload(payload).is_valid


def test_t09_line_id_duplicado(tmp_path: Path) -> None:
    payload, _ = _payload(tmp_path, product_count=2)
    duplicate = replace(payload.products[1], line_id=payload.products[0].line_id)
    invalid = replace(payload, products=(payload.products[0], duplicate))
    assert "line_id_duplicado" in {item.code for item in validate_payload(invalid).errors}
    with pytest.raises(DocumentValidationError):
        generate_word(invalid, tmp_path.resolve())


@pytest.mark.parametrize(("case", "mode"), (("T10", "generated"), ("T11", "manual"), ("T12", "mixed")))
def test_t10_to_t12_origenes_coste(case: str, mode: str, tmp_path: Path) -> None:
    payload, _ = _payload(tmp_path, product_count=5, mode=mode)
    origins = {product.cost_origin for product in payload.products}
    if mode == "generated":
        assert origins == {"generado"}, case
    elif mode == "manual":
        assert origins == {"manual"}, case
    else:
        assert origins == {"generado", "manual"}, case
    assert validate_excel(generate_excel(payload, tmp_path.resolve()).path, payload).is_valid


def test_t13_productos_bloqueados(tmp_path: Path) -> None:
    payload, _ = _payload(tmp_path, product_count=5, locked=True)
    result = generate_excel(payload, tmp_path.resolve())
    assert all(product.locked for product in payload.products)
    assert validate_excel(result.path, payload).is_valid


@pytest.mark.parametrize(("case", "raw", "display"), (("T14", "0.01", "0,01 €"), ("T15", "-0.01", "-0,01 €")))
def test_t14_t15_residual_visible(case: str, raw: str, display: str, tmp_path: Path) -> None:
    payload, _ = _payload(tmp_path, product_count=5)
    summary = replace(payload.summary, visual_product_residual=RawDisplayValue(raw, display))
    changed = replace(payload, summary=summary)
    result = generate_excel(changed, tmp_path.resolve())
    assert validate_excel(result.path, changed).is_valid, case


@pytest.mark.parametrize(("case", "image_format"), (("T16", "PNG"), ("T17", "JPEG")))
def test_t16_t17_imagen_valida(case: str, image_format: str, tmp_path: Path) -> None:
    payload, image = _payload(tmp_path, image_format=image_format)
    assert inspect_route_image(image).mime_type == f"image/{'jpeg' if image_format == 'JPEG' else 'png'}"
    assert not generate_word(payload, tmp_path.resolve(), route_image_path=image).warnings, case
    assert not generate_excel(payload, tmp_path.resolve(), route_image_path=image).warnings, case


def test_t18_imagen_ausente(tmp_path: Path) -> None:
    payload, _ = _payload(tmp_path)
    word = generate_word(payload, tmp_path.resolve())
    excel = generate_excel(payload, tmp_path.resolve())
    assert word.warnings == excel.warnings == ("imagen_ruta_ausente",)


def test_t19_imagen_corrupta(tmp_path: Path) -> None:
    path = tmp_path / "corrupt.png"
    path.write_bytes(b"not-an-image")
    with pytest.raises(InvalidRouteImageError):
        inspect_route_image(path)


def test_t20_imagen_extrema(tmp_path: Path) -> None:
    path = _route_image(tmp_path / "extreme.png", size=(10_001, 1))
    with pytest.raises(InvalidRouteImageError):
        inspect_route_image(path)


def test_t21_marcador_obligatorio_ausente(tmp_path: Path) -> None:
    with zipfile.ZipFile(DEFAULT_TEMPLATE_PATH) as archive:
        data = archive.read("word/document.xml").replace(b"{{ narrative.subject }}", b"ASUNTO")
    changed = tmp_path / "missing.docx"
    _rewrite_zip(DEFAULT_TEMPLATE_PATH, changed, {"word/document.xml": data})
    assert "marcadores_obligatorios_ausentes" in {item.code for item in validate_template(changed).errors}


def test_t22_marcador_desconocido(tmp_path: Path) -> None:
    with zipfile.ZipFile(DEFAULT_TEMPLATE_PATH) as archive:
        data = archive.read("word/document.xml").replace(
            WORD_TEMPLATE_SENTINEL.encode(),
            f"{WORD_TEMPLATE_SENTINEL} {{{{ unknown_value }}}}".encode(),
        )
    changed = tmp_path / "unknown.docx"
    _rewrite_zip(DEFAULT_TEMPLATE_PATH, changed, {"word/document.xml": data})
    assert "marcadores_desconocidos" in {item.code for item in validate_template(changed).errors}


def test_t23_marcador_dividido_entre_runs(tmp_path: Path) -> None:
    with zipfile.ZipFile(DEFAULT_TEMPLATE_PATH) as archive:
        data = archive.read("word/document.xml").replace(
            b"{{ narrative.subject }}",
            b'{{ narrative.</w:t></w:r><w:r><w:t xml:space="preserve">subject }}',
        )
    changed = tmp_path / "split.docx"
    _rewrite_zip(DEFAULT_TEMPLATE_PATH, changed, {"word/document.xml": data})
    assert "marcador_dividido_entre_runs" in {item.code for item in validate_template(changed).errors}


def test_t24_marcador_en_tabla() -> None:
    with zipfile.ZipFile(DEFAULT_TEMPLATE_PATH) as archive:
        xml = archive.read("word/document.xml").decode("utf-8")
    assert "{{ product.name }}" in xml
    assert validate_template(DEFAULT_TEMPLATE_PATH).is_valid


def test_t25_marcadores_en_cabecera_y_pie() -> None:
    with zipfile.ZipFile(DEFAULT_TEMPLATE_PATH) as archive:
        stories = "".join(
            archive.read(name).decode("utf-8")
            for name in archive.namelist()
            if re.fullmatch(r"word/(?:header|footer)\d+\.xml", name)
        )
    assert "{{ identification.expediente }}" in stories
    assert "{{ control.template_version }}" in stories


def test_t26_bloque_opcional_presente_y_ausente(tmp_path: Path) -> None:
    payload, _ = _payload(tmp_path)
    without = generate_word(payload, tmp_path.resolve())
    summary = replace(payload.summary, indirect_costs=RawDisplayValue("100", "100,00 €"))
    with_cost = generate_word(replace(payload, summary=summary), tmp_path.resolve())
    assert "COSTES INDIRECTOS" not in "\n".join(_word_semantics(without.path)[0])
    assert "COSTES INDIRECTOS" in "\n".join(cell for table in _word_semantics(with_cost.path)[1] for row in table for cell in row)


def test_t27_bucle_productos(tmp_path: Path) -> None:
    payload, _ = _payload(tmp_path, product_count=5)
    result = generate_word(payload, tmp_path.resolve())
    document = Document(result.path)
    table = next(table for table in document.tables if table.cell(0, 0).text == "Producto")
    assert len(table.rows) == 6


def test_t28_ausencia_campos_word(tmp_path: Path) -> None:
    payload, _ = _payload(tmp_path)
    result = generate_word(payload, tmp_path.resolve())
    with zipfile.ZipFile(result.path) as archive:
        xml = b"".join(archive.read(name) for name in archive.namelist() if name.startswith("word/") and name.endswith(".xml"))
    assert b"instrText" not in xml and b"fldSimple" not in xml


def test_t29_cabecera_repetible(tmp_path: Path) -> None:
    payload, _ = _payload(tmp_path, product_count=60)
    result = generate_word(payload, tmp_path.resolve())
    document = Document(result.path)
    table = next(table for table in document.tables if table.cell(0, 0).text == "Producto")
    assert table.rows[0]._tr.get_or_add_trPr().find(qn("w:tblHeader")) is not None


def test_t30_filas_sin_altura_fija(tmp_path: Path) -> None:
    payload, _ = _payload(tmp_path, product_count=5, long=True)
    result = generate_word(payload, tmp_path.resolve())
    document = Document(result.path)
    for table in document.tables:
        for row in table.rows:
            height = row._tr.get_or_add_trPr().find(qn("w:trHeight"))
            assert height is None or height.get(qn("w:hRule")) != "exact"


def test_t31_excel_sin_formulas(tmp_path: Path) -> None:
    payload, _ = _payload(tmp_path)
    result = generate_excel(payload, tmp_path.resolve())
    workbook = load_workbook(result.path, data_only=False)
    assert not [cell.coordinate for sheet in workbook for row in sheet.iter_rows() for cell in row if cell.data_type == "f"]
    workbook.close()


def test_t32_snapshot_hash_incorrecto(tmp_path: Path) -> None:
    payload, _ = _payload(tmp_path)
    report = validate_payload(payload, expected_snapshot_sha256="0" * 64)
    assert "snapshot_hash_no_coincide" in {item.code for item in report.errors}


def test_t33_version_incompatible(tmp_path: Path) -> None:
    payload, _ = _payload(tmp_path)
    invalid = replace(payload, control=replace(payload.control, payload_schema_version="999"))
    assert "payload_version_incompatible" in {item.code for item in validate_payload(invalid).errors}


def test_t34_metadatos_y_rutas_personales(tmp_path: Path) -> None:
    payload, _ = _payload(tmp_path)
    word = generate_word(payload, tmp_path.resolve())
    with zipfile.ZipFile(word.path) as archive:
        root = ET.fromstring(archive.read("docProps/core.xml"))
        creator = root.find("{http://purl.org/dc/elements/1.1/}creator")
        assert creator is not None
        creator.text = r"C:\Users\Personal"
        core = ET.tostring(root, encoding="utf-8", xml_declaration=True)
    tainted = tmp_path / "tainted.docx"
    _rewrite_zip(word.path, tainted, {"docProps/core.xml": core})
    codes = {item.code for item in validate_word(tainted, payload).errors}
    assert {"metadatos_word", "ruta_personal_word"} <= codes


def test_t35_enlaces_externos_excel(tmp_path: Path) -> None:
    payload, _ = _payload(tmp_path)
    excel = generate_excel(payload, tmp_path.resolve())
    relation = b'''<?xml version="1.0" encoding="UTF-8"?><Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"><Relationship Id="rId1" Type="x" Target="file:///C:/Users/Personal/source.xlsx" TargetMode="External"/></Relationships>'''
    tainted = tmp_path / "external.xlsx"
    _rewrite_zip(excel.path, tainted, {}, {"xl/externalLinks/_rels/externalLink1.xml.rels": relation})
    codes = {item.code for item in validate_excel(tainted, payload).errors}
    assert "enlaces_externos_excel" in codes


def test_t36_regeneracion_semanticamente_estable(tmp_path: Path) -> None:
    payload, _ = _payload(tmp_path, product_count=5)
    word1 = generate_word(payload, tmp_path.resolve())
    word2 = generate_word(payload, tmp_path.resolve())
    excel1 = generate_excel(payload, tmp_path.resolve())
    excel2 = generate_excel(payload, tmp_path.resolve())
    assert _word_semantics(word1.path) == _word_semantics(word2.path)
    assert _excel_semantics(excel1.path) == _excel_semantics(excel2.path)


def test_t37_proteccion_sobrescritura(tmp_path: Path) -> None:
    payload, _ = _payload(tmp_path)
    first = generate_word(payload, tmp_path.resolve(), version=1)
    with pytest.raises(FileExistsError):
        generate_word(payload, tmp_path.resolve(), version=1)
    assert first.path.exists()


def test_t38_referencias_originales_intactas() -> None:
    assert {name: _sha256(REFERENCE_DIR / name) for name in REFERENCE_HASHES} == REFERENCE_HASHES


def test_t39_unicode_y_caracteres_espanoles(tmp_path: Path) -> None:
    identity = IdentificationInput(
        expediente="EXP-Ñ-ÁRBOL-2026",
        organismo="Diputación Pública",
        objeto="Suministro de alimentación y café",
        lot_number="Lote nº 1",
        lot_name="Panadería",
        duration_description="doce meses",
        place="Córdoba",
        date_text="14 de julio de 2026",
        client="COMPAÑÍA DE PRUEBA, S.L.",
    )
    payload, _ = _payload(tmp_path, product_count=1, identification=identity)
    result = generate_word(payload, tmp_path.resolve())
    assert "Ñ" in result.path.name and validate_word(result.path, payload).is_valid


def test_t40_valores_opcionales_nulos(tmp_path: Path) -> None:
    identity = IdentificationInput(
        expediente="EXP-NULL-1",
        organismo="Organismo de prueba",
        objeto="Objeto de prueba",
        lot_number="1",
        lot_name="",
        duration_description="",
        place="",
        date_text="",
        client="Cliente de prueba",
    )
    payload, _ = _payload(tmp_path, product_count=1, identification=identity)
    assert validate_payload(payload).is_valid
    assert validate_word(generate_word(payload, tmp_path.resolve()).path, payload).is_valid


def test_escritura_atomica_y_no_sobrescritura(tmp_path: Path) -> None:
    final = tmp_path / "atomic.docx"
    temp = temporary_output_path(final)
    temp.write_bytes(b"complete")
    publish_atomic_no_overwrite(temp, final)
    assert final.read_bytes() == b"complete" and not temp.exists()
    another = temporary_output_path(final)
    another.write_bytes(b"replacement")
    with pytest.raises(FileExistsError):
        publish_atomic_no_overwrite(another, final)
    assert final.read_bytes() == b"complete"


def test_nombres_seguros_y_rutas_peligrosas(tmp_path: Path) -> None:
    assert safe_component("Expediente Ñ / 2026") == "Expediente_Ñ_2026"
    with pytest.raises(UnsafeDocumentPathError):
        safe_component("../escape")
    with pytest.raises(UnsafeDocumentPathError):
        next_versioned_path("relative", prefix="x", expediente="e", lot_number="1", suffix=".docx")


def test_generadores_no_importan_motor_calculo() -> None:
    source_root = DEFAULT_TEMPLATE_PATH.parent.parent
    for filename in ("word_generator.py", "excel_generator.py"):
        tree = ast.parse((source_root / filename).read_text(encoding="utf-8"))
        imported = {
            alias.name
            for node in ast.walk(tree)
            if isinstance(node, (ast.Import, ast.ImportFrom))
            for alias in node.names
        }
        assert "calculate_justification" not in imported
        assert "calculations" not in imported


def test_igualdad_word_excel_payload(tmp_path: Path) -> None:
    payload, image = _payload(tmp_path, image_format="PNG")
    word = generate_word(payload, tmp_path.resolve(), route_image_path=image)
    excel = generate_excel(payload, tmp_path.resolve(), route_image_path=image)
    assert validate_word(word.path, payload).is_valid
    assert validate_excel(excel.path, payload).is_valid
    workbook = load_workbook(excel.path, data_only=False)
    audit = workbook["Auditoría"]
    values = {audit.cell(row, 1).value: audit.cell(row, 2).value for row in range(1, audit.max_row + 1)}
    workbook.close()
    assert values["snapshot_sha256"] == payload.control.snapshot_sha256
    assert values["payload_sha256"] == payload.sha256


def test_plantilla_productiva_normalizada() -> None:
    assert validate_template(DEFAULT_TEMPLATE_PATH).is_valid
    with zipfile.ZipFile(DEFAULT_TEMPLATE_PATH) as archive:
        names = archive.namelist()
        xml = b"".join(archive.read(name) for name in names if name.startswith("word/") and name.endswith(".xml"))
    assert WORD_TEMPLATE_SENTINEL.encode() in xml
    assert b"<w:sdt" not in xml
    assert b"instrText" not in xml and b"fldSimple" not in xml
    assert not any(name.startswith("word/comments") for name in names)
