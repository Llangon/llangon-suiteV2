from __future__ import annotations

import hashlib
import io
import json
import zipfile
from datetime import datetime, timezone
from functools import lru_cache
from pathlib import Path
from xml.etree import ElementTree as ET

import pytest
from docx import Document

from herramientas_python.descargadores.common.document_model import (
    build_question_document,
    format_document_question_heading,
)

from herramientas_python.descargadores.common.docx_renderer import (
    CONTENT_WIDTH_DXA,
    DARK_GRAY,
    DOCX_OUTPUT,
    LABEL_WIDTH_DXA,
    TABLE_INDENT_DXA,
    VALUE_WIDTH_DXA,
    render_question_document_docx,
    validate_docx_content,
)
from herramientas_python.descargadores.common.question_models import (
    DocumentRenderError,
    PlatformQuestion,
    SafeFileError,
)
from herramientas_python.descargadores.common.question_state import (
    _empty_state,
    state_directory,
    state_file,
    transaction_file,
)
from herramientas_python.descargadores.common.question_workflow import (
    record_successful_review,
    recover_pending_transaction,
    regenerate_document_from_state,
    write_new_questions_rtf,
)
from herramientas_python.descargadores.common.rtf_renderer import (
    render_question_document,
    validate_rtf_content,
)
from herramientas_python.descargadores.common.safe_files import (
    atomic_write_json,
    publish_document,
)
from webapp.infonalia_webapp.tests.place_docx_visual_scenarios import (
    GENERATED_AT,
    metadata,
    visual_scenarios,
)


W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
CP_NS = "http://schemas.openxmlformats.org/package/2006/metadata/core-properties"
DC_NS = "http://purl.org/dc/elements/1.1/"
EP_NS = "http://schemas.openxmlformats.org/officeDocument/2006/extended-properties"


def normal_model():
    return visual_scenarios()["01_normal"]


@lru_cache(maxsize=1)
def rendered_normal() -> bytes:
    return render_question_document_docx(normal_model())


@lru_cache(maxsize=None)
def rendered_scenario(scenario_name: str) -> bytes:
    return render_question_document_docx(visual_scenarios()[scenario_name])


def rendered(model=None) -> bytes:
    return render_question_document_docx(model) if model is not None else rendered_normal()


def package_parts(content: bytes) -> dict[str, bytes]:
    with zipfile.ZipFile(io.BytesIO(content)) as package:
        return {name: package.read(name) for name in package.namelist()}


def xml_root(content: bytes, name: str):
    return ET.fromstring(package_parts(content)[name])


def xml_text(content: bytes, name: str) -> str:
    return package_parts(content)[name].decode("utf-8", errors="ignore")


def visible_text(content: bytes) -> str:
    root = xml_root(content, "word/document.xml")
    return "\n".join(
        "".join(node.text or "" for node in paragraph.iter(f"{{{W_NS}}}t"))
        for paragraph in root.iter(f"{{{W_NS}}}p")
    )


def test_document_model_keeps_semantic_and_heading_dates_separate() -> None:
    model = dated_document_model(
        asked_at="2026-07-08T10:28:00Z",
        answered_at="2026-07-08T12:40:00Z",
        official_datetime="2026-07-08T10:28:00Z",
    )
    question = model.questions[0]

    assert question.asked_at == "2026-07-08T10:28:00Z"
    assert question.answered_at == "2026-07-08T12:40:00Z"
    assert question.heading_at == "2026-07-08T10:28:00Z"
    assert format_document_question_heading(question) == "Pregunta 3 del 08-07-2026 a las 12:28"


def test_docx_uses_answer_date_as_neutral_heading_without_duplicate_note() -> None:
    model = dated_document_model(
        platform="CATALUNYA",
        number=27,
        asked_at="",
        answered_at="2026-07-17T12:40:08.727Z",
        official_datetime="2026-07-17T12:40:08.727Z",
    )
    question = model.questions[0]
    text = visible_text(rendered(model))

    assert question.asked_at == ""
    assert question.answered_at == "2026-07-17T12:40:08.727Z"
    assert question.heading_at == "2026-07-17T12:40:08.727Z"
    assert "Pregunta 27 del 17-07-2026 a las 14:40" in text
    assert "Respuesta publicada el 17-07-2026 a las 14:40." not in text


def test_docx_keeps_distinct_answer_publication_date_for_place() -> None:
    model = dated_document_model(
        asked_at="2026-07-08T10:28:00Z",
        answered_at="2026-07-08T12:40:00Z",
        official_datetime="2026-07-08T10:28:00Z",
    )
    text = visible_text(rendered(model))

    assert "Pregunta 3 del 08-07-2026 a las 12:28" in text
    assert "Respuesta publicada el 08-07-2026 a las 14:40." in text


@pytest.mark.parametrize(
    ("answered_at", "expected"),
    (
        ("2026-07-17T12:40:08.727Z", "Pregunta 27 del 17-07-2026 a las 14:40"),
        ("2026-01-17T12:40:08Z", "Pregunta 27 del 17-01-2026 a las 13:40"),
    ),
)
def test_docx_heading_uses_madrid_summer_and_winter_time(answered_at: str, expected: str) -> None:
    model = dated_document_model(
        platform="CATALUNYA",
        number=27,
        answered_at=answered_at,
        official_datetime=answered_at,
    )

    assert expected in visible_text(rendered(model))


def test_document_heading_without_any_reliable_date_does_not_invent_one() -> None:
    model = dated_document_model(
        number=9,
        asked_at="fecha desconocida",
        answered_at="sin fecha",
        official_datetime="",
    )
    question = model.questions[0]
    text = visible_text(rendered(model))

    assert question.heading_at == ""
    assert format_document_question_heading(question) == "Pregunta 9"
    assert "Pregunta 9 del" not in text
    assert "Pregunta 9" in text


def platform_question(*, text: str = "¿Se admite el formato A?", answer: str = "Sí."):
    return PlatformQuestion(
        updated_at="08-07-2026 12:28",
        question=text,
        answer=answer,
        source_id="Q-17",
        platform="PLACE",
    )


def dated_document_model(
    *,
    platform: str = "PLACE",
    number: int = 3,
    asked_at: str = "",
    answered_at: str = "",
    official_datetime: str = "",
    timezone_name: str = "Europe/Madrid",
    published: bool = True,
    publication_history=(),
    attachments=(),
    versions=None,
):
    question_text = "Pregunta literal"
    answer_text = "Respuesta literal"
    version_items = versions or [
        {
            "version": 1,
            "detected_at": "2026-07-18T10:00:00+02:00",
            "question": question_text,
            "answer": answer_text,
            "attachments": list(attachments),
            "change_type": "initial",
            "changed_fields": [],
            "asked_at": asked_at,
            "answered_at": answered_at,
        }
    ]
    stored = {
        "stable_id": f"{platform.lower()}-{number}",
        "number": number,
        "official_datetime": official_datetime,
        "asked_at": asked_at,
        "answered_at": answered_at,
        "question": question_text,
        "answer": answer_text,
        "attachments": list(attachments),
        "published": published,
        "publication_history": list(publication_history),
        "versions": version_items,
    }
    tender = metadata()
    tender.update({"platform": platform, "display_timezone": timezone_name})
    return build_question_document(
        tender,
        {"platform": platform, "questions": {stored["stable_id"]: stored}},
        GENERATED_AT,
    )


def test_docx_renderer_returns_zip_bytes() -> None:
    content = rendered()
    assert isinstance(content, bytes)
    assert content.startswith(b"PK")
    assert zipfile.is_zipfile(io.BytesIO(content))


def test_docx_package_contains_required_parts_and_standard_content_type() -> None:
    parts = package_parts(rendered())
    assert {
        "[Content_Types].xml",
        "_rels/.rels",
        "docProps/core.xml",
        "docProps/app.xml",
        "word/document.xml",
        "word/styles.xml",
        "word/_rels/document.xml.rels",
    } <= set(parts)
    assert b"wordprocessingml.document.main+xml" in parts["[Content_Types].xml"]


def test_docx_opens_with_python_docx() -> None:
    document = Document(io.BytesIO(rendered()))
    assert document.paragraphs
    assert document.tables


def test_docx_uses_a4_portrait_and_explicit_margins() -> None:
    root = xml_root(rendered(), "word/document.xml")
    page_size = root.find(f".//{{{W_NS}}}pgSz")
    margins = root.find(f".//{{{W_NS}}}pgMar")
    assert page_size is not None and margins is not None
    assert page_size.attrib[f"{{{W_NS}}}w"] == "11906"
    assert page_size.attrib[f"{{{W_NS}}}h"] == "16838"
    assert page_size.attrib.get(f"{{{W_NS}}}orient", "portrait") == "portrait"
    assert margins.attrib[f"{{{W_NS}}}left"] == "1134"
    assert margins.attrib[f"{{{W_NS}}}right"] == "1134"
    assert margins.attrib[f"{{{W_NS}}}top"] == "850"
    assert margins.attrib[f"{{{W_NS}}}bottom"] == "907"


def test_docx_corporate_header_and_title_are_real_paragraphs() -> None:
    text = visible_text(rendered())
    for expected in (
        "ASESORES LLANGON, S.L.",
        "CIF B73803637",
        "C/ ULIA, 9, 1.º D, 41005, SEVILLA",
        "617 11 02 81 · info@llangon.com",
        "PREGUNTAS Y RESPUESTAS",
        "Documento actualizado el 18/07/2026 a las 16:30",
    ):
        assert expected in text


def test_docx_tender_table_has_only_requested_main_fields() -> None:
    text = visible_text(rendered())
    for label in ("Expediente", "Órgano de contratación", "Objeto", "Fin de presentación", "Enlace"):
        assert label in text
    for removed in (
        "Estado:",
        "Tipo de contrato",
        "Procedimiento",
        "Presupuesto sin IVA",
        "Valor estimado",
        "CPV",
        "Lugar de ejecución",
    ):
        assert removed not in text


def test_docx_tender_table_uses_fixed_dxa_geometry() -> None:
    root = xml_root(rendered(), "word/document.xml")
    table = root.find(f".//{{{W_NS}}}tbl")
    assert table is not None
    table_width = table.find(f"./{{{W_NS}}}tblPr/{{{W_NS}}}tblW")
    table_indent = table.find(f"./{{{W_NS}}}tblPr/{{{W_NS}}}tblInd")
    grid = table.findall(f"./{{{W_NS}}}tblGrid/{{{W_NS}}}gridCol")
    assert table_width.attrib[f"{{{W_NS}}}w"] == str(CONTENT_WIDTH_DXA)
    assert table_indent.attrib[f"{{{W_NS}}}w"] == str(TABLE_INDENT_DXA)
    assert [int(item.attrib[f"{{{W_NS}}}w"]) for item in grid] == [
        LABEL_WIDTH_DXA,
        VALUE_WIDTH_DXA,
    ]
    assert all(
        int(cell.attrib[f"{{{W_NS}}}w"]) in {LABEL_WIDTH_DXA, VALUE_WIDTH_DXA}
        for cell in table.findall(f".//{{{W_NS}}}tcW")
    )


def test_docx_link_is_real_short_external_hyperlink() -> None:
    content = rendered()
    text = visible_text(content)
    document_root = xml_root(content, "word/document.xml")
    rels_root = xml_root(content, "word/_rels/document.xml.rels")
    hyperlinks = list(document_root.iter(f"{{{W_NS}}}hyperlink"))
    relations = {
        relation.attrib["Id"]: relation
        for relation in rels_root.findall(f"{{{REL_NS}}}Relationship")
    }
    assert text.count("Abrir ficha de la licitación en PLACE") == 1
    assert "https://example.test/licitacion/EXP-17" not in text
    assert len(hyperlinks) == 1
    relation = relations[hyperlinks[0].attrib[f"{{{R_NS}}}id"]]
    assert relation.attrib["Target"] == "https://example.test/licitacion/EXP-17"
    assert relation.attrib["TargetMode"] == "External"


def test_docx_hyperlink_is_dark_gray_underlined_and_not_theme_colored() -> None:
    root = xml_root(rendered(), "word/document.xml")
    hyperlink = next(root.iter(f"{{{W_NS}}}hyperlink"))
    color = hyperlink.find(f".//{{{W_NS}}}color")
    underline = hyperlink.find(f".//{{{W_NS}}}u")
    assert color is not None and color.attrib[f"{{{W_NS}}}val"] == DARK_GRAY
    assert f"{{{W_NS}}}themeColor" not in color.attrib
    assert underline is not None and underline.attrib[f"{{{W_NS}}}val"] == "single"


def test_docx_visible_document_is_monochrome_and_has_no_fills() -> None:
    document_xml = xml_text(rendered(), "word/document.xml")
    assert all(value not in document_xml for value in ("2E74B5", "0563C1", "954F72", "0000FF"))
    assert "<w:shd" not in document_xml
    assert "w:highlight" not in document_xml


def test_docx_question_order_is_descending_with_stable_tie_break() -> None:
    text = visible_text(rendered_scenario("02_varias_ordenadas"))
    assert text.index("Pregunta 2 del 08-07-2026 a las 12:28") < text.index(
        "Pregunta 3 del 08-07-2026 a las 10:21"
    )
    assert text.index("Pregunta 3 del 08-07-2026 a las 10:21") < text.index(
        "Pregunta 1 del 03-07-2026 a las 14:14"
    )


def test_docx_styles_encode_numeric_font_spacing_and_keep_rules() -> None:
    styles = xml_root(rendered(), "word/styles.xml")
    question_style = next(
        style
        for style in styles.findall(f"{{{W_NS}}}style")
        if style.attrib.get(f"{{{W_NS}}}styleId") == "LlangonQuestionHeading"
    )
    size = question_style.find(f"./{{{W_NS}}}rPr/{{{W_NS}}}sz")
    spacing = question_style.find(f"./{{{W_NS}}}pPr/{{{W_NS}}}spacing")
    keep_next = question_style.find(f"./{{{W_NS}}}pPr/{{{W_NS}}}keepNext")
    assert size is not None and size.attrib[f"{{{W_NS}}}val"] == "22"
    assert spacing is not None and spacing.attrib[f"{{{W_NS}}}before"] == "220"
    assert spacing.attrib[f"{{{W_NS}}}after"] == "100"
    assert keep_next is not None


def test_docx_uses_keep_with_next_without_forcing_long_body_together() -> None:
    styles = xml_root(rendered_scenario("12_contenido_largo"), "word/styles.xml")
    by_id = {
        style.attrib.get(f"{{{W_NS}}}styleId"): style
        for style in styles.findall(f"{{{W_NS}}}style")
    }
    assert by_id["LlangonLabel"].find(f"./{{{W_NS}}}pPr/{{{W_NS}}}keepNext") is not None
    for property_name in ("keepNext", "keepLines"):
        node = by_id["LlangonBody"].find(f"./{{{W_NS}}}pPr/{{{W_NS}}}{property_name}")
        assert node is None or node.attrib.get(f"{{{W_NS}}}val") in {"0", "false", "off"}


def test_docx_core_properties_are_corporate_and_path_free() -> None:
    content = rendered()
    core = xml_root(content, "docProps/core.xml")
    assert core.findtext(f"{{{DC_NS}}}creator") == "Llangon"
    assert core.findtext(f"{{{CP_NS}}}lastModifiedBy") == "Llangon"
    assert core.findtext(f"{{{DC_NS}}}title") == "PREGUNTAS Y RESPUESTAS"
    assert core.findtext(f"{{{DC_NS}}}subject") == "Expediente EXP-17/2026"
    raw = ET.tostring(core, encoding="unicode")
    assert "C:\\" not in raw and "Users\\" not in raw


def test_docx_extended_properties_identify_company() -> None:
    app = xml_root(rendered(), "docProps/app.xml")
    assert app.findtext(f"{{{EP_NS}}}Company") == "Llangon"


def test_docx_package_has_no_local_paths_temporary_names_or_credentials() -> None:
    parts = package_parts(rendered())
    package_text = "\n".join(
        payload.decode("utf-8", errors="ignore")
        for name, payload in parts.items()
        if name.endswith((".xml", ".rels"))
    ).casefold()
    for forbidden in (
        "c:\\users\\",
        "file:///",
        ".document-",
        "place_contrasena",
        "place_password",
        "configured-secret",
    ):
        assert forbidden not in package_text


def test_docx_package_has_no_macros_embeddings_or_altchunks() -> None:
    parts = package_parts(rendered())
    names = "\n".join(parts).casefold()
    xml = "\n".join(
        payload.decode("utf-8", errors="ignore")
        for name, payload in parts.items()
        if name.endswith((".xml", ".rels"))
    )
    assert "vbaproject" not in names
    assert "activex" not in names
    assert "embeddings/" not in names
    assert "macroEnabled" not in xml
    assert "<w:altChunk" not in xml


def test_docx_attachment_is_preserved_as_real_hyperlink() -> None:
    content = rendered_scenario("11_adjunto")
    text = visible_text(content)
    rels = xml_root(content, "word/_rels/document.xml.rels")
    targets = [
        relation.attrib.get("Target")
        for relation in rels.findall(f"{{{REL_NS}}}Relationship")
        if relation.attrib.get("TargetMode") == "External"
    ]
    assert "Archivos adjuntos" in text
    assert "Aclaración técnica del órgano de contratación.pdf" in text
    assert "Ref. ATT-17" in text
    assert "https://example.test/adjuntos/aclaracion.pdf" in targets


def test_docx_notices_and_all_versions_are_visible() -> None:
    text = visible_text(rendered_scenario("10_historial_extenso"))
    assert "AVISO: CONTENIDO MODIFICADO EN PLACE" in text
    assert "HISTORIAL DE VERSIONES" in text
    for expected in ("Pregunta inicial", "Pregunta revisada", "Pregunta vigente"):
        assert expected in text
    for expected in ("Respuesta inicial", "Respuesta revisada", "Respuesta vigente"):
        assert expected in text


@pytest.mark.parametrize(
    "scenario_name",
    ("04_pregunta_modificada", "08_pregunta_retirada", "09_pregunta_reaparecida", "11_adjunto"),
)
def test_versions_withdrawals_restorations_and_attachments_keep_common_heading(
    scenario_name: str,
) -> None:
    text = visible_text(rendered_scenario(scenario_name))

    assert "Pregunta 1 del 08-07-2026 a las 12:28" in text


def test_docx_empty_answer_uses_explicit_platform_message() -> None:
    text = visible_text(rendered_scenario("03_pendiente"))
    assert "Sin respuesta publicada en PLACE." in text


def test_docx_long_content_is_complete_and_not_truncated() -> None:
    model = visual_scenarios()["12_contenido_largo"]
    text = visible_text(rendered_scenario("12_contenido_largo"))
    assert model.questions[0].versions[0].question_text in text
    assert model.questions[0].versions[0].answer_text in text


def test_docx_preserves_unicode_braces_backslashes_and_line_breaks() -> None:
    content = rendered_scenario("13_caracteres_y_sin_fecha")
    text = visible_text(content)
    assert "Pregunta 1 del" not in text
    assert "Pregunta 1" in text
    assert "¿Se admite A/B, piñón, café y símbolos { } \\?" in text
    assert "Segunda línea." in text
    validate_docx_content(content)


def test_docx_invalid_package_is_rejected() -> None:
    with pytest.raises(DocumentRenderError):
        validate_docx_content(b"no-es-un-docx")


def test_docx_local_file_link_is_not_embedded() -> None:
    model = visual_scenarios()["01_normal"]
    fields = tuple(
        type(field)(field.label, field.value, type(field.link)("file:///C:/Users/demo/secret", field.link.label))
        if field.link
        else field
        for field in model.tender_fields
    )
    local_model = type(model)(
        corporate=model.corporate,
        title=model.title,
        updated_text=model.updated_text,
        tender_section_title=model.tender_section_title,
        tender_fields=fields,
        questions=model.questions,
        generated_at=model.generated_at,
        preferred_page_breaks=model.preferred_page_breaks,
    )
    content = rendered(local_model)
    assert b"file:///" not in package_parts(content)["word/_rels/document.xml.rels"]


@pytest.mark.parametrize("scenario_name", sorted(visual_scenarios()))
def test_all_thirteen_visual_scenarios_generate_valid_docx(scenario_name: str) -> None:
    assert len(visual_scenarios()) == 13
    content = rendered_scenario(scenario_name)
    validate_docx_content(content)
    assert len(content) > 30_000


def test_official_flow_creates_only_docx_and_generic_result(tmp_path: Path) -> None:
    result = record_successful_review(
        tmp_path,
        metadata(),
        [platform_question()],
        reviewed_at=GENERATED_AT,
    )
    assert result.document_generated is True
    assert result.document_format == result.generated_format == "docx"
    assert result.document_name.endswith(".docx")
    assert Path(result.document_path).is_file()
    assert result.document_sha256 == hashlib.sha256(Path(result.document_path).read_bytes()).hexdigest()
    assert result.rtf_generated is False and result.rtf_path == ""
    assert len(list(tmp_path.glob("Preguntas y respuestas a fecha *.docx"))) == 1
    assert not list(tmp_path.glob("Preguntas y respuestas a fecha *.rtf"))


def test_official_no_change_review_creates_no_new_document(tmp_path: Path) -> None:
    record_successful_review(tmp_path, metadata(), [platform_question()], reviewed_at=GENERATED_AT)
    repeated = record_successful_review(
        tmp_path,
        metadata(),
        [platform_question()],
        reviewed_at=datetime(2026, 7, 18, 17, 0, tzinfo=timezone.utc),
    )
    assert repeated.status == "no_changes"
    assert repeated.document_generated is False
    assert repeated.document_path == repeated.document_format == repeated.document_sha256 == ""
    assert repeated.rtf_generated is False and repeated.rtf_path == ""
    assert len(list(tmp_path.glob("Preguntas y respuestas a fecha *.docx"))) == 1


def test_state_last_result_is_generic_and_legacy_rtf_fields_are_empty(tmp_path: Path) -> None:
    record_successful_review(tmp_path, metadata(), [platform_question()], reviewed_at=GENERATED_AT)
    state = json.loads(state_file(tmp_path).read_text(encoding="utf-8"))
    result = state["last_result"]
    assert result["document_generated"] is True
    assert result["document_format"] == "docx"
    assert result["document_name"].endswith(".docx")
    assert len(result["document_sha256"]) == 64
    assert result["rtf_generated"] is False and result["rtf_path"] == ""
    assert result["errors"] == []


def test_explicit_regeneration_creates_docx_without_changing_state(tmp_path: Path) -> None:
    first = record_successful_review(
        tmp_path,
        metadata(),
        [platform_question()],
        reviewed_at=GENERATED_AT,
    )
    before_state = state_file(tmp_path).read_bytes()
    before_document = Path(first.document_path).read_bytes()
    regenerated = regenerate_document_from_state(
        tmp_path,
        generated_at=datetime(2026, 7, 18, 18, 0, tzinfo=timezone.utc),
    )
    assert regenerated.status == "regenerated"
    assert regenerated.document_generated is True
    assert regenerated.no_changes is True and regenerated.changes_detected is False
    assert state_file(tmp_path).read_bytes() == before_state
    assert Path(first.document_path).read_bytes() == before_document
    assert len(list(tmp_path.glob("Preguntas y respuestas a fecha *.docx"))) == 2


def test_explicit_regeneration_resolves_filename_collision_without_overwrite(tmp_path: Path) -> None:
    first = record_successful_review(
        tmp_path,
        metadata(),
        [platform_question()],
        reviewed_at=GENERATED_AT,
    )
    first_hash = hashlib.sha256(Path(first.document_path).read_bytes()).hexdigest()
    regenerated = regenerate_document_from_state(tmp_path, generated_at=GENERATED_AT)
    assert Path(regenerated.document_path).name.endswith("16-30-01.docx")
    assert hashlib.sha256(Path(first.document_path).read_bytes()).hexdigest() == first_hash


def test_binary_transaction_recovery_validates_docx_and_commits_state(tmp_path: Path) -> None:
    technical = state_directory(tmp_path, create=True)
    target = tmp_path / "Preguntas y respuestas a fecha 2026-07-18 16-30-00.docx"
    target.write_bytes(rendered())
    state = _empty_state("https://example.test", metadata(), platform="PLACE")
    atomic_write_json(
        transaction_file(tmp_path),
        {
            "target_name": target.name,
            "temporary_name": "",
            "document_format": "docx",
            "document_extension": ".docx",
            "state": state,
        },
    )
    assert recover_pending_transaction(tmp_path) is True
    assert json.loads(state_file(tmp_path).read_text(encoding="utf-8"))["schema_version"] == 2
    assert not transaction_file(tmp_path).exists()
    assert technical.is_dir()


def test_binary_publication_never_overwrites_and_cleans_failed_temporary(tmp_path: Path) -> None:
    content = rendered()
    technical = state_directory(tmp_path, create=True)
    target = tmp_path / "Preguntas y respuestas a fecha 2026-07-18 16-30-00.docx"
    first = publish_document(
        destination=tmp_path,
        technical_directory=technical,
        content=content,
        target=target,
        output=DOCX_OUTPUT,
    )
    with pytest.raises(SafeFileError):
        publish_document(
            destination=tmp_path,
            technical_directory=technical,
            content=content,
            target=target,
            output=DOCX_OUTPUT,
        )
    assert first.path == target
    assert target.read_bytes() == content
    assert not list(technical.glob(".document-*.tmp"))


def test_historical_rtf_renderer_remains_functional_but_is_not_official(tmp_path: Path) -> None:
    model = normal_model()
    rtf = render_question_document(model)
    validate_rtf_content(rtf)
    path, count = write_new_questions_rtf(
        tmp_path,
        metadata(),
        [platform_question()],
        generated_at=GENERATED_AT,
    )
    assert rtf.startswith(r"{\rtf1")
    assert path is not None and path.suffix == ".rtf" and path.is_file()
    assert count == 1
    assert not list(tmp_path.glob("Preguntas y respuestas a fecha *.docx"))
