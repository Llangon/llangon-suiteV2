"""Pure Word generation from DocumentPayloadV1."""

from __future__ import annotations

import hashlib
from decimal import Decimal
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Mm, Pt
from docxtpl import DocxTemplate, InlineImage
from jinja2 import Environment, StrictUndefined

from .filenames import (
    next_versioned_path,
    publish_atomic_no_overwrite,
    temporary_output_path,
)
from .generation import DocumentGenerationResult
from .metadata import clear_word_core_properties, scrub_docx_package
from .ooxml_helpers import (
    set_cell_margins,
    set_repeat_table_header,
    set_row_cant_split,
    set_table_geometry,
)
from .payload import DocumentPayloadV1
from .template_manifest import DEFAULT_TEMPLATE_PATH, WORD_TEMPLATE_VERSION
from .validators import (
    DocumentValidationError,
    InvalidRouteImageError,
    inspect_route_image,
    validate_payload,
    validate_template,
    validate_word,
)


PRODUCT_HEADERS = (
    "Producto",
    "Características",
    "Cantidad",
    "Precio oferta",
    "Importe oferta",
    "Precio coste",
    "Importe coste",
)
PRODUCT_WIDTHS_DXA = (1800, 2800, 950, 1100, 1300, 1100, 1150)


def generate_word(
    payload: DocumentPayloadV1,
    output_directory: str | Path,
    *,
    route_image_path: str | Path | None = None,
    template_path: str | Path = DEFAULT_TEMPLATE_PATH,
    version: int | None = None,
) -> DocumentGenerationResult:
    """Render and atomically publish one editable Word document."""

    validate_payload(payload).require_valid()
    if payload.control.template_version != WORD_TEMPLATE_VERSION:
        raise ValueError("El payload no corresponde a la plantilla Word productiva.")
    template = Path(template_path)
    validate_template(template).require_valid()
    image_path, warnings = _validated_image(payload, route_image_path)
    final_path, selected_version = next_versioned_path(
        output_directory,
        prefix="Justificacion_Baja",
        expediente=payload.identification.expediente,
        lot_number=payload.identification.lot_number,
        suffix=".docx",
        version=version,
    )
    raw_temp = temporary_output_path(final_path)
    processed_temp = temporary_output_path(final_path)
    scrubbed_temp = temporary_output_path(final_path)
    try:
        template_document = DocxTemplate(str(template))
        context = payload.to_dict()
        context["products_empty"] = not payload.products
        context["show_indirect_costs"] = Decimal(payload.summary.indirect_costs.raw) != 0
        if image_path is not None:
            context["route_image"] = InlineImage(
                template_document,
                str(image_path),
                width=Mm(150),
            )
            context["route_image_note"] = ""
        else:
            context["route_image"] = ""
            context["route_image_note"] = "[IMAGEN DE RUTA PENDIENTE]"
        template_document.render(
            context,
            jinja_env=Environment(undefined=StrictUndefined),
            autoescape=True,
        )
        template_document.save(raw_temp)
        _postprocess_word(raw_temp, processed_temp)
        scrub_docx_package(processed_temp, scrubbed_temp)
        report = validate_word(scrubbed_temp, payload)
        if not report.is_valid:
            raise DocumentValidationError(report)
        publish_atomic_no_overwrite(scrubbed_temp, final_path)
        final_report = validate_word(final_path, payload)
        if not final_report.is_valid:
            final_path.unlink(missing_ok=True)
            raise DocumentValidationError(final_report)
        return DocumentGenerationResult(
            path=final_path,
            sha256=_sha256_file(final_path),
            size_bytes=final_path.stat().st_size,
            warnings=tuple(warnings),
            template_version=payload.control.template_version,
            snapshot_sha256=payload.control.snapshot_sha256,
            payload_sha256=payload.sha256,
            version=selected_version,
        )
    finally:
        for temporary in (raw_temp, processed_temp, scrubbed_temp):
            temporary.unlink(missing_ok=True)


def _validated_image(
    payload: DocumentPayloadV1,
    route_image_path: str | Path | None,
) -> tuple[Path | None, list[str]]:
    reference = payload.transport.route_image
    if route_image_path is None:
        if reference is not None:
            raise InvalidRouteImageError("El payload declara imagen, pero no se proporcionó el archivo.")
        return None, ["imagen_ruta_ausente"]
    if reference is None:
        raise InvalidRouteImageError("Se proporcionó una imagen no declarada en el payload.")
    path = Path(route_image_path)
    actual = inspect_route_image(path, logical_name=reference.logical_name)
    if (
        actual.logical_name,
        actual.mime_type,
        actual.width_px,
        actual.height_px,
        actual.sha256,
        actual.size_bytes,
    ) != (
        reference.logical_name,
        reference.mime_type,
        reference.width_px,
        reference.height_px,
        reference.sha256,
        reference.size_bytes,
    ):
        raise InvalidRouteImageError("La imagen no coincide con la referencia lógica del payload.")
    return path, []


def _postprocess_word(source: Path, destination: Path) -> None:
    document = Document(source)
    product_table = None
    for table in document.tables:
        if table.rows and tuple(_cell_text(cell) for cell in table.rows[0].cells) == PRODUCT_HEADERS:
            product_table = table
            break
    if product_table is None:
        raise ValueError("La plantilla renderizada no contiene la tabla de productos.")
    set_table_geometry(product_table, PRODUCT_WIDTHS_DXA)
    set_repeat_table_header(product_table.rows[0])
    for table in document.tables:
        for row in table.rows:
            set_row_cant_split(row)
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                set_cell_margins(cell)
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.paragraph_format.line_spacing = 1.0
    for row in product_table.rows[1:]:
        for index, cell in enumerate(row.cells):
            for paragraph in cell.paragraphs:
                paragraph.alignment = (
                    WD_ALIGN_PARAGRAPH.LEFT if index in (0, 1) else WD_ALIGN_PARAGRAPH.RIGHT
                )
    for section in document.sections:
        for story in (section.header, section.footer, section.first_page_header, section.first_page_footer):
            for paragraph in story.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
    clear_word_core_properties(document)
    document.save(destination)


def _cell_text(cell: object) -> str:
    return "\n".join(paragraph.text.strip() for paragraph in cell.paragraphs).strip()


def _sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for chunk in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest().upper()


__all__ = ("generate_word",)
