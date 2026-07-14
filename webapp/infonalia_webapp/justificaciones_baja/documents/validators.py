"""Structural, economic and privacy validators for generated documents."""

from __future__ import annotations

import hashlib
import re
import zipfile
from dataclasses import dataclass
from decimal import Decimal, InvalidOperation
from pathlib import Path
from typing import Iterable
from xml.etree import ElementTree as ET

from docx import Document
from docx.oxml.ns import qn
from docxtpl import DocxTemplate
from openpyxl import load_workbook
from PIL import Image, UnidentifiedImageError

from .payload import DocumentPayloadV1, RouteImageReference
from .template_manifest import (
    DEFAULT_TEMPLATE_MANIFEST,
    PAYLOAD_SCHEMA_VERSION,
    TemplateManifest,
    WORD_TEMPLATE_SENTINEL,
    WORD_TEMPLATE_VERSION,
)


W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
CP_NS = "http://schemas.openxmlformats.org/package/2006/metadata/core-properties"
DC_NS = "http://purl.org/dc/elements/1.1/"
EP_NS = "http://schemas.openxmlformats.org/officeDocument/2006/extended-properties"
JINJA_TOKEN = re.compile(r"({{.*?}}|{%.*?%}|{#.*?#})", re.DOTALL)
HASH_PATTERN = re.compile(r"^[0-9A-F]{64}$")
PERSONAL_PATH_PATTERN = re.compile(
    r"(?:[A-Za-z]:\\Users\\|file:///|\\\\[^\\]+\\|/home/|/Users/)",
    re.IGNORECASE,
)
MAX_IMAGE_BYTES = 10 * 1024 * 1024
MAX_IMAGE_DIMENSION = 10_000
MAX_IMAGE_PIXELS = 40_000_000

PRODUCT_HEADERS = (
    "Producto",
    "Características",
    "Cantidad",
    "Precio oferta",
    "Importe oferta",
    "Precio coste",
    "Importe coste",
)
EXPECTED_EXCEL_SHEETS = (
    "Identificación",
    "Productos",
    "Transporte",
    "Resumen",
    "Auditoría",
)
EXPECTED_WORD_TABLE_COUNT = 6
TRANSPORT_RAW_FIELDS = (
    "circular_kilometres",
    "effective_decimal_hours",
    "kilometre_rate",
    "hourly_rate",
    "temporal_cost",
    "kilometre_cost",
    "full_route_cost",
    "narrative_percentage",
    "allocated_transport",
)


class DocumentValidationError(ValueError):
    """Raised when generation would publish an invalid document."""

    def __init__(self, report: "DocumentValidationReport") -> None:
        self.report = report
        message = "; ".join(f"{issue.code}: {issue.message}" for issue in report.errors)
        super().__init__(message or "La validación documental ha fallado.")


class InvalidRouteImageError(ValueError):
    """Raised when an optional route image is supplied but is unsafe/invalid."""


@dataclass(frozen=True, slots=True)
class DocumentValidationFinding:
    code: str
    message: str
    severity: str = "error"
    location: str | None = None


@dataclass(frozen=True, slots=True)
class DocumentValidationReport:
    findings: tuple[DocumentValidationFinding, ...] = ()

    @property
    def errors(self) -> tuple[DocumentValidationFinding, ...]:
        return tuple(item for item in self.findings if item.severity == "error")

    @property
    def warnings(self) -> tuple[DocumentValidationFinding, ...]:
        return tuple(item for item in self.findings if item.severity != "error")

    @property
    def is_valid(self) -> bool:
        return not self.errors

    def require_valid(self) -> None:
        if not self.is_valid:
            raise DocumentValidationError(self)


def inspect_route_image(
    path: str | Path,
    *,
    logical_name: str | None = None,
) -> RouteImageReference:
    source = Path(path)
    if not source.is_file():
        raise InvalidRouteImageError("La imagen de ruta no existe.")
    size_bytes = source.stat().st_size
    if size_bytes <= 0 or size_bytes > MAX_IMAGE_BYTES:
        raise InvalidRouteImageError("La imagen de ruta supera el tamaño permitido.")
    try:
        with Image.open(source) as image:
            image_format = (image.format or "").upper()
            width, height = image.size
            image.verify()
    except (UnidentifiedImageError, OSError, ValueError) as exc:
        raise InvalidRouteImageError("La imagen de ruta está corrupta o no es reconocible.") from exc
    if image_format not in {"PNG", "JPEG"}:
        raise InvalidRouteImageError("Solo se admiten imágenes PNG o JPEG.")
    if (
        width < 1
        or height < 1
        or width > MAX_IMAGE_DIMENSION
        or height > MAX_IMAGE_DIMENSION
        or width * height > MAX_IMAGE_PIXELS
    ):
        raise InvalidRouteImageError("Las dimensiones de la imagen no son razonables.")
    digest = _sha256_file(source)
    return RouteImageReference(
        logical_name=logical_name or source.name,
        mime_type="image/png" if image_format == "PNG" else "image/jpeg",
        width_px=width,
        height_px=height,
        sha256=digest,
        size_bytes=size_bytes,
    )


def validate_payload(
    payload: DocumentPayloadV1,
    *,
    expected_snapshot_sha256: str | None = None,
) -> DocumentValidationReport:
    findings: list[DocumentValidationFinding] = []
    control = payload.control
    if control.payload_schema_version != PAYLOAD_SCHEMA_VERSION:
        findings.append(_finding("payload_version_incompatible", "Versión de payload incompatible."))
    if control.template_version != WORD_TEMPLATE_VERSION:
        findings.append(_finding("template_version_incompatible", "Versión de plantilla incompatible."))
    if not HASH_PATTERN.fullmatch(control.snapshot_sha256):
        findings.append(_finding("snapshot_hash_invalido", "El hash del snapshot no es SHA-256 canónico."))
    if expected_snapshot_sha256 and control.snapshot_sha256 != expected_snapshot_sha256.upper():
        findings.append(_finding("snapshot_hash_no_coincide", "El hash del snapshot no coincide con el esperado."))
    if not control.snapshot_schema_version or not control.calculation_algorithm_version:
        findings.append(_finding("version_economica_ausente", "Faltan versiones económicas de origen."))
    if not control.generated_at:
        findings.append(_finding("fecha_generacion_ausente", "Falta la fecha de generación."))

    for field_name in ("expediente", "organismo", "objeto", "lot_number", "client"):
        if not getattr(payload.identification, field_name).strip():
            findings.append(
                _finding(
                    "identificacion_obligatoria_ausente",
                    f"Falta el campo obligatorio {field_name}.",
                    location=f"identification.{field_name}",
                )
            )
    if not payload.narrative.estimated_draft_notice.strip():
        findings.append(_finding("marca_borrador_ausente", "Falta la advertencia de borrador."))

    line_ids = [product.line_id for product in payload.products]
    if any(not line_id for line_id in line_ids):
        findings.append(_finding("line_id_ausente", "Todos los productos requieren line_id."))
    duplicates = sorted({line_id for line_id in line_ids if line_ids.count(line_id) > 1})
    if duplicates:
        findings.append(
            _finding(
                "line_id_duplicado",
                f"Hay line_id duplicados: {', '.join(duplicates)}.",
            )
        )
    for product in payload.products:
        if product.cost_origin not in {"generado", "manual", "sin_generar"}:
            findings.append(
                _finding(
                    "origen_coste_invalido",
                    "El origen del coste no está permitido.",
                    location=product.line_id,
                )
            )
        for name, value in (
            ("quantity", product.quantity),
            ("offered_unit_price", product.offered_unit_price),
            ("offered_amount", product.offered_amount),
            ("effective_unit_cost", product.effective_unit_cost),
            ("cost_amount", product.cost_amount),
            ("margin", product.margin),
        ):
            _validate_raw_display(value.raw, value.display, findings, f"{product.line_id}.{name}")
    for name in payload.summary.__dataclass_fields__:
        value = getattr(payload.summary, name)
        _validate_raw_display(value.raw, value.display, findings, f"summary.{name}")
    if payload.summary.allocated_transport.raw != payload.transport.allocated_transport.raw:
        findings.append(_finding("transporte_incoherente", "El transporte del resumen difiere del bloque de transporte."))
    route_image = payload.transport.route_image
    if route_image is not None:
        if not HASH_PATTERN.fullmatch(route_image.sha256):
            findings.append(_finding("imagen_hash_invalido", "El hash lógico de imagen no es válido."))
        if route_image.mime_type not in {"image/png", "image/jpeg"}:
            findings.append(_finding("imagen_mime_invalido", "El tipo lógico de imagen no está permitido."))
    return DocumentValidationReport(tuple(findings))


def validate_template(
    template_path: str | Path,
    manifest: TemplateManifest = DEFAULT_TEMPLATE_MANIFEST,
) -> DocumentValidationReport:
    path = Path(template_path)
    findings: list[DocumentValidationFinding] = []
    if not path.is_file():
        return DocumentValidationReport((_finding("plantilla_ausente", "No existe la plantilla Word."),))
    try:
        template = DocxTemplate(path)
        variables = set(template.get_undeclared_template_variables(context={}))
    except Exception as exc:
        return DocumentValidationReport((_finding("plantilla_invalida", f"La plantilla no puede abrirse: {exc}"),))
    missing = sorted(manifest.required_variables - variables)
    unknown = sorted(variables - manifest.allowed_variables)
    if missing:
        findings.append(_finding("marcadores_obligatorios_ausentes", f"Faltan marcadores: {', '.join(missing)}."))
    if unknown:
        findings.append(_finding("marcadores_desconocidos", f"Hay marcadores desconocidos: {', '.join(unknown)}."))
    with zipfile.ZipFile(path) as archive:
        story_names = _story_xml_names(archive.namelist())
        story_data = [archive.read(name) for name in story_names]
        package_text = "\n".join(data.decode("utf-8", "ignore") for data in story_data)
        if manifest.sentinel not in package_text:
            findings.append(_finding("plantilla_version_incompatible", "No aparece el sentinel de versión."))
        if any(name.endswith("vbaProject.bin") for name in archive.namelist()):
            findings.append(_finding("plantilla_con_macros", "La plantilla no puede contener macros."))
        actual_markers: set[str] = set()
        for name, data in zip(story_names, story_data):
            actual_markers.update(_jinja_tokens(data))
            split = _split_jinja_tokens(data)
            if split:
                findings.append(
                    _finding(
                        "marcador_dividido_entre_runs",
                        f"Hay marcadores divididos entre runs: {', '.join(split[:3])}.",
                        location=name,
                    )
                )
            if b"<w:sdt" in data:
                findings.append(_finding("controles_contenido_plantilla", "La plantilla contiene controles de contenido.", location=name))
            if _field_codes(data):
                findings.append(_finding("campos_word_plantilla", "La plantilla contiene campos Word.", location=name))
        missing_markers = sorted(manifest.required_markers - actual_markers)
        unknown_markers = sorted(actual_markers - manifest.required_markers)
        if missing_markers:
            findings.append(
                _finding(
                    "marcadores_obligatorios_ausentes",
                    f"Faltan marcadores exactos: {', '.join(missing_markers[:5])}.",
                )
            )
        if unknown_markers:
            findings.append(
                _finding(
                    "marcadores_desconocidos",
                    f"Hay marcadores exactos desconocidos: {', '.join(unknown_markers[:5])}.",
                )
            )
    return DocumentValidationReport(tuple(findings))


def validate_word(path: str | Path, payload: DocumentPayloadV1) -> DocumentValidationReport:
    source = Path(path)
    findings = list(validate_payload(payload).findings)
    if not source.is_file():
        findings.append(_finding("word_ausente", "No existe el Word generado."))
        return DocumentValidationReport(tuple(findings))
    try:
        document = Document(source)
    except Exception as exc:
        findings.append(_finding("word_invalido", f"El Word no puede abrirse: {exc}"))
        return DocumentValidationReport(tuple(findings))

    if len(document.tables) != EXPECTED_WORD_TABLE_COUNT:
        findings.append(
            _finding(
                "numero_tablas_word",
                f"Se esperaban {EXPECTED_WORD_TABLE_COUNT} tablas nativas y hay {len(document.tables)}.",
            )
        )

    product_table = _find_word_table(document, PRODUCT_HEADERS)
    if product_table is None:
        findings.append(_finding("tabla_productos_ausente", "No existe la tabla nativa de productos."))
    else:
        expected_rows = [
            [
                product.name,
                product.characteristics,
                product.quantity.display,
                f"{product.offered_unit_price.display} €",
                f"{product.offered_amount.display} €",
                f"{product.effective_unit_cost.display} €",
                f"{product.cost_amount.display} €",
            ]
            for product in payload.products
        ]
        actual_rows = [[_cell_text(cell) for cell in row.cells] for row in product_table.rows[1:]]
        if actual_rows != expected_rows:
            findings.append(_finding("productos_word_no_coinciden", "Las filas Word no coinciden con el payload."))
        if len(product_table.rows) - 1 != len(payload.products):
            findings.append(_finding("numero_productos_word", "El número de productos Word no coincide."))
        header_properties = product_table.rows[0]._tr.get_or_add_trPr()
        if header_properties.find(qn("w:tblHeader")) is None:
            findings.append(_finding("cabecera_no_repetible", "La cabecera de productos no es repetible."))
        for row in product_table.rows:
            properties = row._tr.get_or_add_trPr()
            if properties.find(qn("w:cantSplit")) is None:
                findings.append(_finding("fila_divisible", "Hay filas de productos divisibles."))
                break
            height = properties.find(qn("w:trHeight"))
            if height is not None and height.get(qn("w:hRule")) == "exact":
                findings.append(_finding("altura_fila_exacta", "Hay una altura exacta que puede recortar texto."))
                break

    all_text = _word_all_text(document)
    if payload.control.snapshot_sha256 not in all_text:
        findings.append(_finding("snapshot_hash_word_ausente", "El Word no incluye el hash del snapshot."))
    if payload.narrative.estimated_draft_notice not in all_text:
        findings.append(_finding("marca_borrador_word_ausente", "El Word no muestra la marca de borrador."))
    if JINJA_TOKEN.search(all_text):
        findings.append(_finding("marcadores_word_sin_resolver", "El Word contiene marcadores Jinja."))
    expected_summary = {
        "IMPORTE TOTAL DEL CONTRATO": payload.summary.offer.display,
        "COSTES DE ADQUISICIÓN": payload.summary.prorated_product_cost.display,
        "COSTES DE TRANSPORTE": payload.summary.allocated_transport.display,
        "GASTOS GENERALES": payload.summary.general_expenses.display,
        "COSTE TOTAL": payload.summary.total_cost.display,
        "BENEFICIO": payload.summary.profit.display,
    }
    for label, value in expected_summary.items():
        if label not in all_text or value not in all_text:
            findings.append(_finding("resumen_word_no_coincide", f"Falta {label} o su valor."))
    for value in (
        payload.transport.temporal_cost.display,
        payload.transport.kilometre_cost.display,
        payload.transport.full_route_cost.display,
        payload.transport.allocated_transport.display,
    ):
        if value not in all_text:
            findings.append(
                _finding("transporte_word_no_coincide", f"Falta el valor de transporte {value}.")
            )

    with zipfile.ZipFile(source) as archive:
        names = archive.namelist()
        story_names = _story_xml_names(names)
        story_bytes = [archive.read(name) for name in story_names]
        for name, data in zip(story_names, story_bytes):
            if _field_codes(data):
                findings.append(_finding("campos_word_presentes", "El Word contiene campos dependientes de actualización.", location=name))
            if b"<w:sdt" in data:
                findings.append(_finding("controles_word_presentes", "El Word contiene controles de contenido.", location=name))
        if any(name.lower().startswith("word/comments") for name in names):
            findings.append(_finding("comentarios_word_presentes", "El Word contiene comentarios ocultos."))
        if "docProps/custom.xml" in names:
            findings.append(_finding("propiedades_personalizadas_word", "El Word conserva propiedades personalizadas."))
        media_hashes = {
            hashlib.sha256(archive.read(name)).hexdigest().upper()
            for name in names
            if name.startswith("word/media/")
        }
        route_ref = payload.transport.route_image
        if route_ref is not None and route_ref.sha256 not in media_hashes:
            findings.append(_finding("imagen_ruta_word_no_coincide", "La imagen Word no coincide con su hash lógico."))
        if route_ref is None and "[IMAGEN DE RUTA PENDIENTE]" not in all_text:
            findings.append(_finding("aviso_imagen_word_ausente", "Falta la nota editable de imagen ausente."))
        xml_text = "\n".join(archive.read(name).decode("utf-8", "ignore") for name in names if name.endswith((".xml", ".rels")))
        if PERSONAL_PATH_PATTERN.search(xml_text):
            findings.append(_finding("ruta_personal_word", "El Word contiene una ruta local o personal."))
        if re.search(r"\brsid[A-Za-z]*=", xml_text, re.IGNORECASE):
            findings.append(_finding("rsid_word", "El Word contiene identificadores de sesión."))
        _check_word_metadata(archive, findings)
    return DocumentValidationReport(tuple(_deduplicate(findings)))


def validate_excel(path: str | Path, payload: DocumentPayloadV1) -> DocumentValidationReport:
    source = Path(path)
    findings = list(validate_payload(payload).findings)
    if not source.is_file():
        findings.append(_finding("excel_ausente", "No existe el Excel generado."))
        return DocumentValidationReport(tuple(findings))
    try:
        workbook = load_workbook(source, data_only=False, read_only=False)
    except Exception as exc:
        findings.append(_finding("excel_invalido", f"El Excel no puede abrirse: {exc}"))
        return DocumentValidationReport(tuple(findings))
    if tuple(workbook.sheetnames) != EXPECTED_EXCEL_SHEETS:
        findings.append(_finding("hojas_excel_incorrectas", "Las hojas Excel no coinciden con el contrato."))
    formulas: list[str] = []
    personal_paths: list[str] = []
    for sheet in workbook.worksheets:
        for row in sheet.iter_rows():
            for cell in row:
                if cell.data_type == "f" or (isinstance(cell.value, str) and cell.value.startswith("=")):
                    formulas.append(f"{sheet.title}!{cell.coordinate}")
                if isinstance(cell.value, str) and PERSONAL_PATH_PATTERN.search(cell.value):
                    personal_paths.append(f"{sheet.title}!{cell.coordinate}")
    if formulas:
        findings.append(_finding("formulas_excel_presentes", "El Excel contiene fórmulas."))
    if personal_paths:
        findings.append(_finding("rutas_personales_excel", "El Excel contiene rutas personales."))
    if getattr(workbook, "_external_links", []):
        findings.append(_finding("enlaces_externos_excel", "El Excel contiene enlaces externos."))
    if workbook.properties.creator or workbook.properties.lastModifiedBy:
        findings.append(_finding("metadatos_excel", "El Excel conserva autor o última persona."))

    if "Productos" in workbook.sheetnames:
        sheet = workbook["Productos"]
        ids = [sheet.cell(row, 1).value for row in range(3, sheet.max_row + 1)]
        expected_ids = [product.line_id for product in payload.products]
        if ids != expected_ids:
            findings.append(_finding("productos_excel_no_coinciden", "Los line_id del Excel no coinciden."))
        if len(sheet.tables) != 1:
            findings.append(_finding("tabla_excel_ausente", "Productos debe contener una tabla estructurada."))
    if "Auditoría" in workbook.sheetnames:
        audit = workbook["Auditoría"]
        values = {
            str(audit.cell(row, 1).value): (
                "" if audit.cell(row, 2).value is None else str(audit.cell(row, 2).value)
            )
            for row in range(1, audit.max_row + 1)
            if audit.cell(row, 1).value is not None
        }
        if values.get("snapshot_sha256") != payload.control.snapshot_sha256:
            findings.append(_finding("snapshot_hash_excel", "El hash del snapshot no coincide en Excel."))
        if values.get("payload_sha256") != payload.sha256:
            findings.append(_finding("payload_hash_excel", "El hash del payload no coincide en Excel."))
        for name in payload.summary.__dataclass_fields__:
            key = f"summary.{name}"
            if values.get(key) != getattr(payload.summary, name).raw:
                findings.append(_finding("resumen_excel_no_coincide", f"No coincide {key}."))
        for name in TRANSPORT_RAW_FIELDS:
            key = f"transport.{name}"
            if values.get(key) != getattr(payload.transport, name).raw:
                findings.append(_finding("transporte_excel_no_coincide", f"No coincide {key}."))
        for index, product in enumerate(payload.products):
            prefix = f"product.{index:04d}.{product.line_id}"
            expected = {
                f"{prefix}.line_id": product.line_id,
                f"{prefix}.quantity": product.quantity.raw,
                f"{prefix}.offered_unit_price": product.offered_unit_price.raw,
                f"{prefix}.offered_amount": product.offered_amount.raw,
                f"{prefix}.generated_unit_cost": product.generated_unit_cost_raw or "",
                f"{prefix}.manual_unit_cost": product.manual_unit_cost_raw or "",
                f"{prefix}.effective_unit_cost": product.effective_unit_cost.raw,
                f"{prefix}.cost_amount": product.cost_amount.raw,
                f"{prefix}.margin": product.margin.raw,
                f"{prefix}.cost_origin": product.cost_origin,
                f"{prefix}.locked": str(product.locked).lower(),
            }
            for key, expected_value in expected.items():
                if values.get(key, "") != expected_value:
                    findings.append(
                        _finding("producto_excel_no_coincide", f"No coincide {key}.", location=key)
                    )
    if "Transporte" in workbook.sheetnames:
        images = len(workbook["Transporte"]._images)
        if payload.transport.route_image is not None and images != 1:
            findings.append(_finding("imagen_excel_ausente", "El Excel no contiene la imagen de ruta."))
        if payload.transport.route_image is None and images:
            findings.append(_finding("imagen_excel_inesperada", "El Excel contiene una imagen no declarada."))
    with zipfile.ZipFile(source) as archive:
        names = archive.namelist()
        if any("vbaProject" in name for name in names):
            findings.append(_finding("macros_excel", "El Excel contiene macros."))
        if _package_has_external_relationships(archive):
            findings.append(_finding("enlaces_externos_excel", "El Excel contiene relaciones externas."))
        media_hashes = {
            hashlib.sha256(archive.read(name)).hexdigest().upper()
            for name in names
            if name.startswith("xl/media/")
        }
        route_ref = payload.transport.route_image
        if route_ref is not None and route_ref.sha256 not in media_hashes:
            findings.append(_finding("imagen_excel_no_coincide", "La imagen Excel no coincide con su hash lógico."))
        xml_text = "\n".join(archive.read(name).decode("utf-8", "ignore") for name in names if name.endswith((".xml", ".rels")))
        if PERSONAL_PATH_PATTERN.search(xml_text):
            findings.append(_finding("ruta_paquete_excel", "El paquete Excel contiene rutas personales."))
    workbook.close()
    return DocumentValidationReport(tuple(_deduplicate(findings)))


def _package_has_external_relationships(archive: zipfile.ZipFile) -> bool:
    for name in archive.namelist():
        if not name.endswith(".rels"):
            continue
        try:
            root = ET.fromstring(archive.read(name))
        except ET.ParseError:
            continue
        for relationship in root:
            if relationship.attrib.get("TargetMode", "").lower() == "external":
                return True
    return False


def _validate_raw_display(
    raw: str,
    display: str,
    findings: list[DocumentValidationFinding],
    location: str,
) -> None:
    try:
        value = Decimal(raw)
    except (InvalidOperation, TypeError):
        findings.append(_finding("decimal_documental_invalido", "El valor raw no es decimal.", location=location))
        return
    if not value.is_finite():
        findings.append(_finding("decimal_documental_no_finito", "El valor raw no es finito.", location=location))
    if not str(display).strip():
        findings.append(_finding("presentacion_documental_ausente", "Falta el valor de presentación.", location=location))


def _find_word_table(document: object, header: tuple[str, ...]) -> object | None:
    for table in document.tables:
        if not table.rows:
            continue
        candidate = tuple(_cell_text(cell) for cell in table.rows[0].cells)
        if candidate == header:
            return table
    return None


def _cell_text(cell: object) -> str:
    return "\n".join(paragraph.text.strip() for paragraph in cell.paragraphs).strip()


def _word_all_text(document: object) -> str:
    values: list[str] = []
    values.extend(paragraph.text for paragraph in document.paragraphs)
    for table in document.tables:
        for row in table.rows:
            values.extend(_cell_text(cell) for cell in row.cells)
    for section in document.sections:
        for story in (section.header, section.footer, section.first_page_header, section.first_page_footer):
            values.extend(paragraph.text for paragraph in story.paragraphs)
            for table in story.tables:
                for row in table.rows:
                    values.extend(_cell_text(cell) for cell in row.cells)
    return "\n".join(values)


def _story_xml_names(names: Iterable[str]) -> list[str]:
    return sorted(
        name
        for name in names
        if name == "word/document.xml"
        or re.fullmatch(r"word/(?:header|footer)\d+\.xml", name)
    )


def _field_codes(data: bytes) -> tuple[str, ...]:
    try:
        root = ET.fromstring(data)
    except ET.ParseError:
        return ()
    codes = [node.text or "" for node in root.iter(f"{{{W_NS}}}instrText")]
    simple = [node.attrib.get(f"{{{W_NS}}}instr", "") for node in root.iter(f"{{{W_NS}}}fldSimple")]
    return tuple(code.strip() for code in codes + simple if code.strip())


def _split_jinja_tokens(data: bytes) -> list[str]:
    try:
        root = ET.fromstring(data)
    except ET.ParseError:
        return []
    split: list[str] = []
    for paragraph in root.iter(f"{{{W_NS}}}p"):
        text_nodes = [node.text or "" for node in paragraph.iter(f"{{{W_NS}}}t")]
        joined = "".join(text_nodes)
        for token in JINJA_TOKEN.findall(joined):
            if not any(token in node_text for node_text in text_nodes):
                split.append(token[:80])
    return split


def _jinja_tokens(data: bytes) -> set[str]:
    try:
        root = ET.fromstring(data)
    except ET.ParseError:
        return set()
    result: set[str] = set()
    for paragraph in root.iter(f"{{{W_NS}}}p"):
        joined = "".join(node.text or "" for node in paragraph.iter(f"{{{W_NS}}}t"))
        result.update(" ".join(token.split()) for token in JINJA_TOKEN.findall(joined))
    return result


def _check_word_metadata(
    archive: zipfile.ZipFile,
    findings: list[DocumentValidationFinding],
) -> None:
    names = set(archive.namelist())
    if "docProps/core.xml" in names:
        root = ET.fromstring(archive.read("docProps/core.xml"))
        creator = root.find(f"{{{DC_NS}}}creator")
        modified = root.find(f"{{{CP_NS}}}lastModifiedBy")
        if (creator is not None and (creator.text or "").strip()) or (
            modified is not None and (modified.text or "").strip()
        ):
            findings.append(_finding("metadatos_word", "El Word conserva autor o última persona."))
    if "docProps/app.xml" in names:
        root = ET.fromstring(archive.read("docProps/app.xml"))
        company = root.find(f"{{{EP_NS}}}Company")
        manager = root.find(f"{{{EP_NS}}}Manager")
        if (company is not None and (company.text or "").strip()) or (
            manager is not None and (manager.text or "").strip()
        ):
            findings.append(_finding("empresa_metadatos_word", "El Word conserva empresa o responsable en metadatos."))


def _finding(
    code: str,
    message: str,
    *,
    severity: str = "error",
    location: str | None = None,
) -> DocumentValidationFinding:
    return DocumentValidationFinding(code=code, message=message, severity=severity, location=location)


def _deduplicate(findings: Iterable[DocumentValidationFinding]) -> list[DocumentValidationFinding]:
    result: list[DocumentValidationFinding] = []
    seen: set[tuple[str, str | None]] = set()
    for finding in findings:
        key = (finding.code, finding.location)
        if key not in seen:
            result.append(finding)
            seen.add(key)
    return result


def _sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for chunk in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest().upper()


__all__ = (
    "DocumentValidationError",
    "DocumentValidationFinding",
    "DocumentValidationReport",
    "InvalidRouteImageError",
    "inspect_route_image",
    "validate_excel",
    "validate_payload",
    "validate_template",
    "validate_word",
)
