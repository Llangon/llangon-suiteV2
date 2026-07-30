"""Renderizador DOCX nativo para el modelo documental neutral de preguntas."""

from __future__ import annotations

import io
import re
import zipfile
from datetime import datetime
from urllib.parse import urlsplit
from xml.etree import ElementTree as ET

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Mm, Pt, RGBColor

from .document_model import (
    DocumentAttachment,
    DocumentNotice,
    DocumentQuestion,
    DocumentVersion,
    QuestionDocument,
    format_document_question_heading,
)
from .question_models import DocumentRenderError, literal_text
from .safe_files import BinaryDocumentOutput


# Preset resuelto: compact_reference_guide / contract_negotiation_brief.
# Overrides reutilizables: A4, sistema monocromo, 10 pt y ficha de licitación A4.
INK = "000000"
DARK_GRAY = "3F3F3F"
MUTED_GRAY = "5A5A5A"
RULE_GRAY = "D0D0D0"
CONTENT_WIDTH_DXA = 9518
TABLE_INDENT_DXA = 120
LABEL_WIDTH_DXA = 2400
VALUE_WIDTH_DXA = CONTENT_WIDTH_DXA - LABEL_WIDTH_DXA
CELL_MARGIN_DXA = 120
ALLOWED_LINK_SCHEMES = {"http", "https"}


def _rgb(value: str) -> RGBColor:
    return RGBColor.from_string(value)


def _set_font(element, name: str) -> None:
    element.font.name = name
    r_pr = element._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    for attribute in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{attribute}"), name)


def _configure_style(
    document,
    name: str,
    *,
    style_type=WD_STYLE_TYPE.PARAGRAPH,
    font_name: str = "Calibri",
    size: float = 10,
    color: str = INK,
    bold: bool = False,
    italic: bool = False,
    before: float = 0,
    after: float = 4,
    line_spacing: float = 1.15,
    keep_with_next: bool = False,
    keep_together: bool = False,
):
    styles = document.styles
    try:
        style = styles[name]
    except KeyError:
        style = styles.add_style(name, style_type)
    _set_font(style, font_name)
    style.font.size = Pt(size)
    style.font.color.rgb = _rgb(color)
    style.font.bold = bold
    style.font.italic = italic
    if style_type == WD_STYLE_TYPE.PARAGRAPH:
        paragraph = style.paragraph_format
        paragraph.space_before = Pt(before)
        paragraph.space_after = Pt(after)
        paragraph.line_spacing = line_spacing
        paragraph.keep_with_next = keep_with_next
        paragraph.keep_together = keep_together
        paragraph.widow_control = True
    return style


def _build_styles(document: QuestionDocument, word_document) -> None:
    corporate = document.corporate
    _configure_style(word_document, "Normal", font_name=corporate.font_name, size=10, after=4)
    _configure_style(
        word_document,
        "Llangon Corporate",
        font_name=corporate.font_name,
        size=11.5,
        bold=True,
        after=2,
        keep_with_next=True,
    )
    _configure_style(
        word_document,
        "Llangon Corporate Detail",
        font_name=corporate.font_name,
        size=8.5,
        color=DARK_GRAY,
        after=1,
        keep_with_next=True,
    )
    _configure_style(
        word_document,
        "Llangon Title",
        font_name=corporate.heading_font_name,
        size=15.5,
        bold=True,
        after=3,
        keep_with_next=True,
    )
    _configure_style(
        word_document,
        "Llangon Updated",
        font_name=corporate.font_name,
        size=8.5,
        color=MUTED_GRAY,
        after=10,
        keep_with_next=True,
    )
    _configure_style(
        word_document,
        "Llangon Section",
        font_name=corporate.font_name,
        size=10.5,
        bold=True,
        before=7,
        after=5,
        keep_with_next=True,
    )
    _configure_style(
        word_document,
        "Llangon Question Heading",
        font_name=corporate.font_name,
        size=11,
        bold=True,
        before=11,
        after=5,
        keep_with_next=True,
    )
    _configure_style(
        word_document,
        "Llangon Label",
        font_name=corporate.font_name,
        size=9,
        bold=True,
        before=2,
        after=2,
        keep_with_next=True,
    )
    _configure_style(
        word_document,
        "Llangon Body",
        font_name=corporate.font_name,
        size=10,
        after=3,
        line_spacing=1.15,
    )
    _configure_style(
        word_document,
        "Llangon Secondary",
        font_name=corporate.font_name,
        size=8.5,
        color=MUTED_GRAY,
        italic=True,
        after=3,
        keep_with_next=True,
    )
    _configure_style(
        word_document,
        "Llangon Notice",
        font_name=corporate.font_name,
        size=9.2,
        bold=True,
        before=4,
        after=3,
        keep_with_next=True,
    )
    _configure_style(
        word_document,
        "Llangon Version",
        font_name=corporate.font_name,
        size=9.5,
        bold=True,
        before=4,
        after=1,
        keep_with_next=True,
    )
    _configure_style(
        word_document,
        "Llangon Attachment",
        font_name=corporate.font_name,
        size=9,
        color=DARK_GRAY,
        after=2,
    )
    _configure_style(
        word_document,
        "Llangon Hyperlink",
        style_type=WD_STYLE_TYPE.CHARACTER,
        font_name=corporate.font_name,
        size=9,
        color=DARK_GRAY,
    ).font.underline = True


def _set_paragraph_border(paragraph, *, top: bool = False, bottom: bool = False) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    existing = p_pr.find(qn("w:pBdr"))
    if existing is not None:
        p_pr.remove(existing)
    borders = OxmlElement("w:pBdr")
    for edge_name, enabled in (("top", top), ("bottom", bottom)):
        if not enabled:
            continue
        edge = OxmlElement(f"w:{edge_name}")
        edge.set(qn("w:val"), "single")
        edge.set(qn("w:sz"), "6")
        edge.set(qn("w:space"), "4")
        edge.set(qn("w:color"), RULE_GRAY)
        borders.append(edge)
    p_pr.append(borders)


def _set_cell_margins(table) -> None:
    tbl_pr = table._tbl.tblPr
    margins = tbl_pr.find(qn("w:tblCellMar"))
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    for side in ("top", "start", "bottom", "end"):
        element = margins.find(qn(f"w:{side}"))
        if element is None:
            element = OxmlElement(f"w:{side}")
            margins.append(element)
        element.set(qn("w:w"), str(80 if side in {"top", "bottom"} else CELL_MARGIN_DXA))
        element.set(qn("w:type"), "dxa")


def _set_table_geometry(table, widths: tuple[int, ...]) -> None:
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    for tag, value in (("tblW", sum(widths)), ("tblInd", TABLE_INDENT_DXA)):
        element = tbl_pr.find(qn(f"w:{tag}"))
        if element is None:
            element = OxmlElement(f"w:{tag}")
            tbl_pr.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)
    for row in table.rows:
        row_properties = row._tr.get_or_add_trPr()
        if row_properties.find(qn("w:cantSplit")) is None:
            row_properties.append(OxmlElement("w:cantSplit"))
        for cell, width in zip(row.cells, widths):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_width = tc_pr.get_or_add_tcW()
            tc_width.set(qn("w:w"), str(width))
            tc_width.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    _set_cell_margins(table)


def _set_quiet_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    existing = tbl_pr.find(qn("w:tblBorders"))
    if existing is not None:
        tbl_pr.remove(existing)
    borders = OxmlElement("w:tblBorders")
    for edge_name in ("top", "left", "bottom", "right"):
        edge = OxmlElement(f"w:{edge_name}")
        edge.set(qn("w:val"), "nil")
        borders.append(edge)
    inside = OxmlElement("w:insideH")
    inside.set(qn("w:val"), "single")
    inside.set(qn("w:sz"), "4")
    inside.set(qn("w:color"), RULE_GRAY)
    borders.append(inside)
    vertical = OxmlElement("w:insideV")
    vertical.set(qn("w:val"), "nil")
    borders.append(vertical)
    tbl_pr.append(borders)


def _safe_link_target(target: object) -> str:
    text = literal_text(target)
    try:
        parsed = urlsplit(text)
    except ValueError:
        return ""
    return text if parsed.scheme.casefold() in ALLOWED_LINK_SCHEMES and parsed.netloc else ""


def _add_hyperlink(paragraph, target: object, label: object) -> bool:
    url = _safe_link_target(target)
    if not url:
        paragraph.add_run(literal_text(label))
        return False
    relationship_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    style = OxmlElement("w:rStyle")
    style.set(qn("w:val"), "LlangonHyperlink")
    run_properties.append(style)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), DARK_GRAY)
    run_properties.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_properties.append(underline)
    run.append(run_properties)
    text = OxmlElement("w:t")
    text.text = literal_text(label)
    run.append(text)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    return True


def _format_cell_paragraph(paragraph, *, label: bool) -> None:
    paragraph.style = "Llangon Label" if label else "Llangon Body"
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.05


def _add_opening_block(word_document, document: QuestionDocument) -> None:
    corporate = document.corporate
    paragraph = word_document.add_paragraph(style="Llangon Corporate")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run(corporate.company_name)
    for text in (
        f"CIF {corporate.tax_id}",
        corporate.address,
        f"{corporate.phone} · {corporate.email}",
    ):
        detail = word_document.add_paragraph(style="Llangon Corporate Detail")
        detail.alignment = WD_ALIGN_PARAGRAPH.CENTER
        detail.add_run(text)
    title = word_document.add_paragraph(style="Llangon Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run(document.title)
    updated = word_document.add_paragraph(style="Llangon Updated")
    updated.alignment = WD_ALIGN_PARAGRAPH.CENTER
    updated.add_run(document.updated_text)


def _add_tender_table(word_document, document: QuestionDocument) -> None:
    heading = word_document.add_paragraph(style="Llangon Section")
    heading.add_run(document.tender_section_title)
    fields = [field for field in document.tender_fields if literal_text(field.value)]
    if not fields:
        return
    table = word_document.add_table(rows=len(fields), cols=2)
    for row, field in zip(table.rows, fields):
        label_paragraph = row.cells[0].paragraphs[0]
        _format_cell_paragraph(label_paragraph, label=True)
        label_paragraph.add_run(field.label)
        value_paragraph = row.cells[1].paragraphs[0]
        _format_cell_paragraph(value_paragraph, label=False)
        if field.link:
            _add_hyperlink(value_paragraph, field.link.target, field.link.label)
        else:
            value_paragraph.add_run(field.value)
    _set_table_geometry(table, (LABEL_WIDTH_DXA, VALUE_WIDTH_DXA))
    _set_quiet_table_borders(table)


def _add_notice(word_document, notice: DocumentNotice):
    title = word_document.add_paragraph(style="Llangon Notice")
    title.add_run(notice.title)
    _set_paragraph_border(title, top=True, bottom=True)
    last = title
    for line in notice.lines:
        if not literal_text(line):
            continue
        last = word_document.add_paragraph(style="Llangon Secondary")
        last.add_run(line)
    return last


def _add_attachments(word_document, attachments: tuple[DocumentAttachment, ...]):
    if not attachments:
        return None
    roles = {literal_text(item.role).casefold() for item in attachments}
    label_text = {
        frozenset({"answer"}): "Archivos adjuntos a la respuesta",
        frozenset({"question"}): "Archivos adjuntos a la pregunta",
    }.get(frozenset(roles), "Archivos adjuntos")
    label = word_document.add_paragraph(style="Llangon Label")
    label.add_run(label_text)
    last = label
    for attachment in attachments:
        last = word_document.add_paragraph(style="Llangon Attachment")
        last.paragraph_format.left_indent = Mm(4)
        if attachment.link:
            _add_hyperlink(last, attachment.link.target, attachment.link.label)
        else:
            last.add_run(attachment.name)
        if attachment.source_id:
            secondary = last.add_run(f"  ·  Ref. {attachment.source_id}")
            secondary.italic = True
            secondary.font.color.rgb = _rgb(MUTED_GRAY)
    return last


def _add_version(word_document, version: DocumentVersion):
    last = None
    if version.label:
        last = word_document.add_paragraph(style="Llangon Version")
        last.add_run(version.label)
    if version.date_note:
        last = word_document.add_paragraph(style="Llangon Secondary")
        last.add_run(version.date_note)
    if version.show_question_label:
        last = word_document.add_paragraph(style="Llangon Label")
        last.add_run("Pregunta")
    if version.question_date_note:
        last = word_document.add_paragraph(style="Llangon Secondary")
        last.add_run(version.question_date_note)
    last = word_document.add_paragraph(style="Llangon Body")
    last.paragraph_format.keep_with_next = True
    last.add_run(version.question_text)
    answer_label = word_document.add_paragraph(style="Llangon Label")
    answer_label.add_run("Respuesta")
    if version.answer_date_note:
        last = word_document.add_paragraph(style="Llangon Secondary")
        last.add_run(version.answer_date_note)
    last = word_document.add_paragraph(style="Llangon Body")
    answer_text = literal_text(version.answer_text)
    answer_run = last.add_run(answer_text or version.empty_answer_text)
    if not answer_text:
        answer_run.italic = True
        answer_run.font.color.rgb = _rgb(MUTED_GRAY)
    if version.attachments:
        last.paragraph_format.keep_with_next = True
    attachment_paragraph = _add_attachments(word_document, version.attachments)
    return attachment_paragraph or last


def _add_question(word_document, question: DocumentQuestion) -> None:
    heading = word_document.add_paragraph(style="Llangon Question Heading")
    heading.add_run(format_document_question_heading(question))
    _set_paragraph_border(heading, bottom=True)
    last = heading
    for notice in question.publication_notices:
        last = _add_notice(word_document, notice)
    if question.modification_notice:
        last = _add_notice(word_document, question.modification_notice)
    if question.show_version_history_heading:
        last = word_document.add_paragraph(style="Llangon Version")
        last.add_run("HISTORIAL DE VERSIONES")
    for version in question.versions:
        last = _add_version(word_document, version)
    _set_paragraph_border(last, bottom=True)
    last.paragraph_format.space_after = Pt(8)


def _set_core_properties(word_document, document: QuestionDocument, generated_at: datetime) -> None:
    expediente = next(
        (field.value for field in document.tender_fields if field.label == "Expediente"),
        "",
    )
    timestamp = generated_at.astimezone().replace(tzinfo=None) if generated_at.tzinfo else generated_at
    properties = word_document.core_properties
    properties.author = "Llangon"
    properties.last_modified_by = "Llangon"
    properties.title = document.title
    properties.subject = f"Expediente {expediente}" if expediente else "Preguntas y respuestas de licitación"
    properties.keywords = "licitación; preguntas y respuestas; plataforma de contratación"
    properties.comments = "Documento corporativo generado por Llangon"
    properties.category = "Licitaciones"
    properties.created = timestamp
    properties.modified = timestamp
    properties.revision = 1


def _patch_company_property(content: bytes) -> bytes:
    source_buffer = io.BytesIO(content)
    output_buffer = io.BytesIO()
    with zipfile.ZipFile(source_buffer, "r") as source, zipfile.ZipFile(
        output_buffer,
        "w",
    ) as output:
        for member in source.infolist():
            payload = source.read(member.filename)
            if member.filename == "docProps/app.xml":
                root = ET.fromstring(payload)
                namespace = "http://schemas.openxmlformats.org/officeDocument/2006/extended-properties"
                company = root.find(f"{{{namespace}}}Company")
                if company is None:
                    company = ET.SubElement(root, f"{{{namespace}}}Company")
                company.text = "Llangon"
                payload = ET.tostring(root, encoding="utf-8", xml_declaration=True)
            output.writestr(member, payload)
    return output_buffer.getvalue()


def render_question_document_docx(document: QuestionDocument) -> bytes:
    """Genera un DOCX directamente desde ``QuestionDocument`` y devuelve sus bytes."""

    word_document = Document()
    section = word_document.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(15)
    section.bottom_margin = Mm(16)
    section.left_margin = Mm(20)
    section.right_margin = Mm(20)
    section.header_distance = Mm(8)
    section.footer_distance = Mm(8)
    _build_styles(document, word_document)
    try:
        generated_at = datetime.fromisoformat(document.generated_at)
    except ValueError as exc:
        raise DocumentRenderError("El modelo documental no contiene una fecha de generación válida.") from exc
    _set_core_properties(word_document, document, generated_at)
    _add_opening_block(word_document, document)
    _add_tender_table(word_document, document)
    questions_heading = word_document.add_paragraph(style="Llangon Section")
    questions_heading.paragraph_format.space_before = Pt(10)
    questions_heading.add_run("PREGUNTAS RESPONDIDAS")
    if document.questions:
        for question in document.questions:
            _add_question(word_document, question)
    else:
        empty = word_document.add_paragraph(style="Llangon Body")
        empty.add_run("No hay preguntas publicadas en el estado consultado.")
    buffer = io.BytesIO()
    word_document.save(buffer)
    content = _patch_company_property(buffer.getvalue())
    validate_docx_content(content)
    return content


def validate_docx_content(content: bytes) -> None:
    """Valida integridad, relaciones, metadatos y seguridad básica del paquete."""

    if not isinstance(content, bytes) or not zipfile.is_zipfile(io.BytesIO(content)):
        raise DocumentRenderError("El DOCX generado no es un paquete ZIP válido.")
    required = {
        "[Content_Types].xml",
        "_rels/.rels",
        "docProps/core.xml",
        "docProps/app.xml",
        "word/document.xml",
        "word/styles.xml",
        "word/_rels/document.xml.rels",
    }
    with zipfile.ZipFile(io.BytesIO(content), "r") as package:
        names = set(package.namelist())
        missing = required - names
        if missing:
            raise DocumentRenderError(
                "El DOCX generado está incompleto: " + ", ".join(sorted(missing))
            )
        lowered_names = {name.casefold() for name in names}
        if any(
            marker in name
            for name in lowered_names
            for marker in ("vbaproject", "activex", "embeddings/")
        ):
            raise DocumentRenderError("El DOCX generado contiene componentes activos no permitidos.")
        package_text = "\n".join(
            package.read(name).decode("utf-8", errors="ignore")
            for name in names
            if name.endswith((".xml", ".rels"))
        )
        if "macroEnabled" in package_text or "<w:altChunk" in package_text:
            raise DocumentRenderError("El DOCX generado contiene contenido activo no permitido.")
        local_path_patterns = (
            r"(?i)[a-z]:\\(?:users|windows|program files)\\",
            r"(?i)file:/+",
            r"(?i)\\\\[^\\\s]+\\",
            r"(?i)\.document-[0-9a-f]{16,}",
        )
        if any(re.search(pattern, package_text) for pattern in local_path_patterns):
            raise DocumentRenderError("El DOCX generado expone una ruta local o temporal.")
        content_types = package.read("[Content_Types].xml").decode("utf-8")
        expected_main = (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"
        )
        if expected_main not in content_types:
            raise DocumentRenderError("El tipo de contenido principal del DOCX no es válido.")
        document_xml = ET.fromstring(package.read("word/document.xml"))
        rels_xml = ET.fromstring(package.read("word/_rels/document.xml.rels"))
        rel_namespace = "http://schemas.openxmlformats.org/package/2006/relationships"
        relationships = {
            relation.attrib.get("Id", ""): relation
            for relation in rels_xml.findall(f"{{{rel_namespace}}}Relationship")
        }
        r_namespace = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
        w_namespace = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        for hyperlink in document_xml.iter(f"{{{w_namespace}}}hyperlink"):
            relationship_id = hyperlink.attrib.get(f"{{{r_namespace}}}id", "")
            relation = relationships.get(relationship_id)
            if relation is None:
                raise DocumentRenderError("El DOCX contiene un hipervínculo sin relación válida.")
            if relation.attrib.get("TargetMode") != "External":
                raise DocumentRenderError("El DOCX contiene un hipervínculo externo mal declarado.")
            if not _safe_link_target(relation.attrib.get("Target")):
                raise DocumentRenderError("El DOCX contiene un destino de hipervínculo no permitido.")
            colors = [
                node.attrib.get(f"{{{w_namespace}}}val", "").upper()
                for node in hyperlink.iter(f"{{{w_namespace}}}color")
            ]
            if not colors or any(color != DARK_GRAY for color in colors):
                raise DocumentRenderError("El DOCX contiene un hipervínculo azul o sin color corporativo.")
        core_text = package.read("docProps/core.xml").decode("utf-8", errors="ignore")
        app_text = package.read("docProps/app.xml").decode("utf-8", errors="ignore")
        if "Llangon" not in core_text or "Llangon" not in app_text:
            raise DocumentRenderError("Los metadatos corporativos del DOCX están incompletos.")
    try:
        Document(io.BytesIO(content))
    except Exception as exc:
        raise DocumentRenderError("Microsoft Word no podrá abrir el DOCX generado.") from exc


DOCX_OUTPUT = BinaryDocumentOutput(
    format_name="docx",
    extension=".docx",
    render=render_question_document_docx,
    validator=validate_docx_content,
)
